import streamlit as st
import anthropic
from PyPDF2 import PdfReader
import json, re, html, io
import plotly.express as px
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from prompt import (
    SYSTEM_PROMPT, GENRE_RULES,
    build_analysis_prompt, build_doctoring_prompt, build_rewrite_prompt,
    get_report_filename
)

# =================================================================
# [0] 페이지 설정 & 세션 초기화
# =================================================================
st.set_page_config(page_title="BLUE JEANS REWRITE ENGINE", layout="wide", page_icon="👖")

for k, v in {
    "page": "index",
    "db": [],
    "selected_item": None,
    "step": 0,
    "raw_text": None,
    "analysis": None,
    "washing": None,
    "rewriting": None,
}.items():
    if k not in st.session_state:
        st.session_state[k] = v

# =================================================================
# [1] 디자인 시스템
# =================================================================
# ── 장르별 Unsplash 이미지 URL 생성 ──
def get_genre_img_url(item):
    """장르별 Picsum Photos 이미지 (CORS 제한 없음, 완전 무료)"""
    genre_raw = item.get('genre', {})
    if isinstance(genre_raw, dict):
        genre_str = genre_raw.get('primary', '') + ' ' + ' '.join(genre_raw.get('tags', []))
    else:
        genre_str = str(genre_raw)
    genre_str = genre_str.lower()

    # 장르별 Picsum 이미지 ID (영화적 분위기 큐레이션)
    # https://picsum.photos/id/{id}/400/200 형태
    genre_ids = {
        'romance':  [1024, 488, 429, 543, 718],   # 따뜻한 색감, 커플, 자연
        'action':   [1036, 373, 412, 365, 307],   # 어둡고 강렬한
        'comedy':   [1082, 247, 669, 835, 593],   # 밝고 경쾌한
        'horror':   [1067, 167, 202, 239, 376],   # 어둡고 미스터리
        'fantasy':  [1019, 325, 462, 533, 672],   # 환상적, 하늘
        'family':   [1060, 219, 342, 491, 610],   # 따뜻한 가족
        'crime':    [1043, 110, 164, 288, 395],   # 도시, 밤
        'history':  [1029, 145, 271, 390, 508],   # 고전적
        'drama':    [1062, 232, 366, 487, 614],   # 감성적
    }

    mapping = [
        (['로맨스','romance','멜로','love','로코'],  'romance'),
        (['액션','action','스릴러','thriller'],      'action'),
        (['코미디','comedy'],                        'comedy'),
        (['공포','horror','호러'],                   'horror'),
        (['sf','sci-fi','판타지','fantasy','타임'],  'fantasy'),
        (['가족','family','성장'],                   'family'),
        (['범죄','crime','누아르','noir'],           'crime'),
        (['전쟁','war','역사','historical','사극'],  'history'),
        (['드라마','drama'],                         'drama'),
    ]

    genre_key = 'drama'
    for keys, gkey in mapping:
        if any(k in genre_str for k in keys):
            genre_key = gkey
            break

    title = item.get('title', 'film')
    ids = genre_ids[genre_key]
    img_id = ids[abs(hash(title)) % len(ids)]
    return f"https://picsum.photos/id/{img_id}/400/200"


def apply_design():
    st.markdown("""
    <style>
    @import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');

    :root { color-scheme: light only !important; }

    .stApp,
    [data-testid="stAppViewContainer"],
    [data-testid="stHeader"],
    [data-testid="stMain"],
    [data-testid="stMainBlockContainer"],
    [data-testid="block-container"],
    body { background-color: #FAFAFA !important; }

    html, body, .stApp, p, span, h1, h2, h3, h4, h5,
    label, li, td, th {
        color: #191970 !important;
        font-family: 'Pretendard', sans-serif !important;
    }
    div {
        font-family: 'Pretendard', sans-serif !important;
    }

    [data-testid="stSidebar"] { display: none !important; }

    /* 배경색별 텍스트 강제 */
    [style*="background:#FFCB05"] { color: #191970 !important; }
    [style*="background:#FFCB05"] * { color: #191970 !important; }
    [style*="background:#2EC484"] { color: #FFFFFF !important; }
    [style*="background:#2EC484"] * { color: #FFFFFF !important; }

    /* ── th 텍스트 색상 보호 (와일드카드 덮어씌움 방지) ── */
    th[style*="color:#FFCB05"],
    th[style*="color: #FFCB05"] {
        color: #FFCB05 !important;
    }
    th[style*="background:#191970"] {
        color: #FFCB05 !important;
        background-color: #191970 !important;
    }
    th[style*="background:#191970"] * {
        color: #FFCB05 !important;
    }

    /* ── Expander arrow 수정 ── */
    /* 헤더 배경 */
    [data-testid="stExpander"] details summary {
        background: #F0F2FF !important;
        border-radius: 8px !important;
        padding: 10px 16px !important;
        border: 1px solid #E0E4F0 !important;
    }
    [data-testid="stExpander"] details[open] summary {
        border-radius: 8px 8px 0 0 !important;
        border-bottom: 1px solid #FFCB05 !important;
    }
    /* 화살표 SVG 색상 명시 */
    [data-testid="stExpander"] summary svg {
        color: #191970 !important;
        fill: #191970 !important;
        stroke: #191970 !important;
        opacity: 1 !important;
    }
    /* 결과 보기 expander (agent card 안) */
    [data-testid="stExpander"] details summary p,
    [data-testid="stExpander"] details summary span {
        color: #191970 !important;
        font-weight: 700 !important;
    }
    /* expander 내용 영역 */
    [data-testid="stExpander"] details > div {
        border: 1px solid #E0E4F0 !important;
        border-top: none !important;
        border-radius: 0 0 8px 8px !important;
        padding: 16px !important;
        background: #FAFAFA !important;
    }

    .agent-card {
        background: #FFFFFF;
        border: 1px solid #E6E9EF;
        border-radius: 16px;
        padding: 24px 28px;
        margin-bottom: 20px;
        position: relative;
        overflow: hidden;
        box-shadow: 0 2px 12px rgba(25,25,112,0.06);
    }
    .agent-card::before {
        content: '';
        position: absolute;
        top: 0; left: 0; right: 0;
        height: 4px;
        background: linear-gradient(90deg, #FFCB05, #191970);
    }
    .agent-card.active { border-color: #FFCB05; box-shadow: 0 4px 20px rgba(255,203,5,0.2); }
    .agent-card.done   { border-color: #B8E6D4; }
    .agent-card.waiting { opacity: 0.42; }

    .report-card {
        background: #FFFFFF !important;
        border: 1px solid #E6E9EF !important;
        border-left: 5px solid #FFCB05 !important;
        border-radius: 12px !important;
        padding: 22px !important;
        margin-bottom: 16px !important;
        box-shadow: 0 2px 8px rgba(25,25,112,0.04) !important;
    }
    .report-card h3 {
        color: #191970 !important;
        font-size: 0.95rem !important;
        font-weight: 800 !important;
        margin-bottom: 14px !important;
        letter-spacing: 0.04em !important;
    }

    /* ── 파일 업로더 버튼: Streamlit 기본 사용 ── */

    /* ── #191970 배경: 흰색 텍스트 강제 ── */
    [style*="background:#191970"],
    [style*="background: #191970"] {
        color: #FFFFFF !important;
    }
    [style*="background:#191970"] span,
    [style*="background:#191970"] div,
    [style*="background: #191970"] span,
    [style*="background: #191970"] div {
        color: #FFFFFF !important;
    }

    /* ── report-card 안 span 인라인 배경 보호 ── */
    .report-card span[style*="background:#191970"],
    div[style*="background:#FFFBE6"] span[style*="background:#191970"] {
        background: #191970 !important;
        background-color: #191970 !important;
        color: #FFCB05 !important;
    }

    /* ── 테이블 헤더 th: 와일드카드 덮어쓰기 방지 ── */
    th[style*="color:#FFCB05"],
    th[style*="color: #FFCB05"] {
        color: #FFCB05 !important;
        background: #191970 !important;
    }
    thead tr[style*="background:#191970"] th,
    thead tr[style*="background: #191970"] th {
        color: #FFCB05 !important;
        background: #191970 !important;
    }

    div.stButton > button {
        background: #FFCB05 !important;
        color: #191970 !important;
        font-weight: 900 !important;
        border-radius: 10px !important;
        border: none !important;
        letter-spacing: 0.04em !important;
    }
    div.stButton > button:hover {
        background: #191970 !important;
        color: #FFFFFF !important;
        border: none !important;
    }

    .step-bar { display:flex; align-items:center; justify-content:center; gap:0; margin:24px 0; }
    .step-dot { width:34px; height:34px; border-radius:50%; display:flex; align-items:center; justify-content:center; font-weight:900; font-size:0.8rem; flex-shrink:0; }
    .step-dot.done   { background:#2EC484; color:#FFFFFF; }
    .step-dot.active { background:#FFCB05; color:#191970; }
    .step-dot.wait   { background:#E6E9EF; color:#AAAAAA; border:1px solid #DDD; }
    .step-line { flex:1; height:2px; background:#E6E9EF; }
    .step-line.done { background:#2EC484; }

    .gallery-card {
        background: #FFFFFF;
        border: 1px solid #E6E9EF;
        border-radius: 14px;
        overflow: hidden;
        text-align: center;
        box-shadow: 0 2px 8px rgba(25,25,112,0.04);
    }
    .gallery-card-img {
        width: 100%;
        height: 140px;
        object-fit: cover;
        display: block;
        background: #E6E9EF;
    }
    .gallery-card-body {
        padding: 16px 18px 18px;
    }

    /* ── 스피너 배경 제거 ── */
    [data-testid="stSpinner"],
    [data-testid="stSpinner"] > div,
    .stSpinner,
    div[class*="StatusWidget"],
    div[class*="stStatusWidget"] {
        background: transparent !important;
        background-color: transparent !important;
        box-shadow: none !important;
        border: none !important;
    }
    [data-testid="stSpinner"] p,
    [data-testid="stSpinner"] span,
    .stSpinner p,
    .stSpinner span {
        color: #191970 !important;
        background: transparent !important;
    }

    hr { border-color: #E6E9EF !important; }
    [data-testid="stAlert"] { background: #F0F4FF !important; border-color: #C5D0F0 !important; }
    [data-testid="stDownloadButton"] button {
        background: #191970 !important;
        color: #FFFFFF !important;
        font-weight: 900 !important;
        border-radius: 10px !important;
    }
    /* ── Spinner: 배경 완전 투명 ── */
    [data-testid="stSpinner"] > div,
    [data-testid="stStatusWidget"] {
        background: transparent !important;
        background-color: transparent !important;
    }
    [data-testid="stSpinner"] p,
    [data-testid="stSpinner"] span,
    [data-testid="stSpinner"] div {
        background: transparent !important;
        background-color: transparent !important;
        color: #191970 !important;
    }

    /* ── 파일 업로더: Streamlit 기본 스타일 사용 ── */
    </style>
    """, unsafe_allow_html=True)


# ── Chris: 안경 쓴 분석가, 짧은 단발머리, 네이비 재킷 ──
CHRIS_SVG = """
<svg width="72" height="72" viewBox="0 0 72 72" fill="none" xmlns="http://www.w3.org/2000/svg">
  <rect width="72" height="72" rx="16" fill="#D6E4F0"/>
  <!-- 몸/어깨 -->
  <path d="M14 72 Q14 54 36 50 Q58 54 58 72Z" fill="#1B2A6B"/>
  <!-- 셔츠 칼라 -->
  <path d="M30 50 L36 58 L42 50" fill="white" stroke="#1B2A6B" stroke-width="1.5"/>
  <!-- 목 -->
  <rect x="31" y="42" width="10" height="10" rx="3" fill="#F5C99A"/>
  <!-- 얼굴 -->
  <ellipse cx="36" cy="30" rx="14" ry="15" fill="#F5C99A" stroke="#2D2D2D" stroke-width="1.8"/>
  <!-- 머리카락: 짧고 단정한 옆가르마 -->
  <path d="M22 26 Q22 12 36 11 Q50 12 50 26 Q46 18 36 18 Q26 18 22 26Z" fill="#2D2D2D"/>
  <path d="M22 26 Q20 20 24 16" stroke="#2D2D2D" stroke-width="3" stroke-linecap="round"/>
  <!-- 눈썹 -->
  <path d="M26 26 Q29 24 32 26" stroke="#2D2D2D" stroke-width="2" stroke-linecap="round" fill="none"/>
  <path d="M40 26 Q43 24 46 26" stroke="#2D2D2D" stroke-width="2" stroke-linecap="round" fill="none"/>
  <!-- 안경 프레임 -->
  <rect x="24" y="27" width="10" height="8" rx="3" fill="none" stroke="#2D2D2D" stroke-width="2"/>
  <rect x="38" y="27" width="10" height="8" rx="3" fill="none" stroke="#2D2D2D" stroke-width="2"/>
  <line x1="34" y1="31" x2="38" y2="31" stroke="#2D2D2D" stroke-width="2"/>
  <line x1="22" y1="29" x2="24" y2="30" stroke="#2D2D2D" stroke-width="2"/>
  <line x1="48" y1="29" x2="50" y2="30" stroke="#2D2D2D" stroke-width="2"/>
  <!-- 눈동자 -->
  <circle cx="29" cy="31" r="2" fill="#2D2D2D"/>
  <circle cx="43" cy="31" r="2" fill="#2D2D2D"/>
  <circle cx="30" cy="30.2" r="0.7" fill="white"/>
  <circle cx="44" cy="30.2" r="0.7" fill="white"/>
  <!-- 코 -->
  <path d="M35 33 Q36 36 37 33" stroke="#C8905A" stroke-width="1.2" fill="none" stroke-linecap="round"/>
  <!-- 입: 진지한 표정 -->
  <path d="M30 40 Q36 43 42 40" stroke="#2D2D2D" stroke-width="1.8" fill="none" stroke-linecap="round"/>
</svg>"""

# ── Shiho: 긴 웨이브 머리, 여성 전문가, 보라색 재킷 ──
SHIHO_SVG = """
<svg width="72" height="72" viewBox="0 0 72 72" fill="none" xmlns="http://www.w3.org/2000/svg">
  <rect width="72" height="72" rx="16" fill="#EDE0FF"/>
  <!-- 긴 머리카락(뒤) -->
  <path d="M18 30 Q15 50 18 68 L28 68 Q24 48 24 30Z" fill="#3D2314"/>
  <path d="M54 30 Q57 50 54 68 L44 68 Q48 48 48 30Z" fill="#3D2314"/>
  <!-- 몸/어깨 -->
  <path d="M14 72 Q14 54 36 50 Q58 54 58 72Z" fill="#6B3FA0"/>
  <!-- 목 -->
  <rect x="31" y="42" width="10" height="10" rx="3" fill="#F5C99A"/>
  <!-- 얼굴 -->
  <ellipse cx="36" cy="30" rx="14" ry="15" fill="#F5C99A" stroke="#2D2D2D" stroke-width="1.8"/>
  <!-- 머리카락(앞) -->
  <path d="M22 28 Q22 12 36 11 Q50 12 50 28 Q47 18 36 18 Q25 18 22 28Z" fill="#3D2314"/>
  <!-- 앞머리 -->
  <path d="M22 22 Q28 15 36 16 Q30 17 26 22Z" fill="#3D2314"/>
  <!-- 눈썹: 곡선형 -->
  <path d="M25 25 Q28 22 33 24" stroke="#3D2314" stroke-width="2" stroke-linecap="round" fill="none"/>
  <path d="M39 24 Q44 22 47 25" stroke="#3D2314" stroke-width="2" stroke-linecap="round" fill="none"/>
  <!-- 눈: 크고 또렷 -->
  <ellipse cx="29" cy="29" rx="4" ry="4.5" fill="white" stroke="#2D2D2D" stroke-width="1.5"/>
  <ellipse cx="43" cy="29" rx="4" ry="4.5" fill="white" stroke="#2D2D2D" stroke-width="1.5"/>
  <circle cx="30" cy="30" r="2.5" fill="#3D2314"/>
  <circle cx="44" cy="30" r="2.5" fill="#3D2314"/>
  <circle cx="30.8" cy="29" r="1" fill="white"/>
  <circle cx="44.8" cy="29" r="1" fill="white"/>
  <!-- 속눈썹 -->
  <path d="M25 26 Q27 24 29 25" stroke="#2D2D2D" stroke-width="1.2" fill="none"/>
  <path d="M43 25 Q45 24 47 26" stroke="#2D2D2D" stroke-width="1.2" fill="none"/>
  <!-- 볼터치 -->
  <ellipse cx="22" cy="33" rx="4" ry="2.5" fill="#FFB3C1" opacity="0.55"/>
  <ellipse cx="50" cy="33" rx="4" ry="2.5" fill="#FFB3C1" opacity="0.55"/>
  <!-- 코 -->
  <path d="M35 33 Q36 36 37 33" stroke="#C8905A" stroke-width="1.2" fill="none" stroke-linecap="round"/>
  <!-- 입: 부드러운 미소 -->
  <path d="M29 40 Q36 45 43 40" stroke="#2D2D2D" stroke-width="1.8" fill="none" stroke-linecap="round"/>
  <path d="M31 40 Q36 43 41 40" fill="#FF8FA3" opacity="0.5"/>
  <!-- 귀걸이 -->
  <circle cx="22" cy="32" r="2.5" fill="#FFCB05" stroke="#2D2D2D" stroke-width="1"/>
  <circle cx="50" cy="32" r="2.5" fill="#FFCB05" stroke="#2D2D2D" stroke-width="1"/>
</svg>"""

# ── Moon: 올백머리, 한국 남성, 감독 느낌, 카키 재킷 ──
MOON_SVG = """
<svg width="72" height="72" viewBox="0 0 72 72" fill="none" xmlns="http://www.w3.org/2000/svg">
  <rect width="72" height="72" rx="16" fill="#DFF2E8"/>
  <!-- 몸/어깨 -->
  <path d="M14 72 Q14 54 36 50 Q58 54 58 72Z" fill="#3D6B4F"/>
  <!-- 터틀넥 -->
  <path d="M30 50 Q36 54 42 50 L42 44 Q36 47 30 44Z" fill="#2D2D2D"/>
  <!-- 목 -->
  <rect x="31" y="42" width="10" height="10" rx="3" fill="#E8B88A"/>
  <!-- 얼굴 -->
  <ellipse cx="36" cy="30" rx="14" ry="15" fill="#E8B88A" stroke="#2D2D2D" stroke-width="1.8"/>
  <!-- 머리카락: 올백 스타일 -->
  <path d="M22 26 Q22 10 36 10 Q50 10 50 26 Q48 16 36 15 Q24 15 22 26Z" fill="#1A1A1A"/>
  <path d="M22 24 Q24 14 36 13" stroke="#1A1A1A" stroke-width="4" stroke-linecap="round"/>
  <!-- 올백 라인 -->
  <path d="M26 14 Q36 11 46 14" stroke="#333" stroke-width="1" fill="none" opacity="0.4"/>
  <!-- 눈썹: 진하고 강한 -->
  <path d="M24 25 Q28 22 33 24" stroke="#1A1A1A" stroke-width="2.5" stroke-linecap="round" fill="none"/>
  <path d="M39 24 Q44 22 48 25" stroke="#1A1A1A" stroke-width="2.5" stroke-linecap="round" fill="none"/>
  <!-- 눈: 날카로운 -->
  <path d="M24 29 Q29 27 34 29 Q29 33 24 29Z" fill="#1A1A1A"/>
  <path d="M38 29 Q43 27 48 29 Q43 33 38 29Z" fill="#1A1A1A"/>
  <circle cx="29" cy="29.5" r="1" fill="white"/>
  <circle cx="43" cy="29.5" r="1" fill="white"/>
  <!-- 코 -->
  <path d="M34 32 Q36 37 38 32" stroke="#B8784A" stroke-width="1.5" fill="none" stroke-linecap="round"/>
  <path d="M32 37 Q36 39 40 37" stroke="#B8784A" stroke-width="1.2" fill="none" stroke-linecap="round"/>
  <!-- 입: 굳은 표정 -->
  <path d="M29 42 Q36 44 43 42" stroke="#2D2D2D" stroke-width="2" fill="none" stroke-linecap="round"/>
  <!-- 인중 -->
  <line x1="36" y1="39" x2="36" y2="42" stroke="#B8784A" stroke-width="1" opacity="0.5"/>
</svg>"""
def safe(text):
    return html.escape(str(text)) if text else ""

def parse_json(raw):
    if not raw: return None

    def try_parse(s):
        try:
            return json.loads(s, strict=False)
        except:
            return None

    def clean(s):
        # BOM 및 유니코드 공백 제거
        s = s.strip().lstrip('\ufeff\u200b\u00a0\u200e\u200f')
        # 마크다운 코드블록 제거
        s = re.sub(r'```json\s*|```\s*', '', s).strip()
        # // 주석 제거
        s = re.sub(r'\s*//[^\n"]*', '', s)
        return s

    # 1차: 그대로
    r = try_parse(raw)
    if r: return r

    # 2차: 클린업 후
    cleaned = clean(raw)
    r = try_parse(cleaned)
    if r: return r

    # 3차: { } 블록 추출 후 클린업
    try:
        start = raw.index('{')
        end   = raw.rindex('}') + 1
        chunk = clean(raw[start:end])
        r = try_parse(chunk)
        if r: return r
    except:
        pass

    # 4차: simplejson으로 재시도 (더 관대한 파서)
    try:
        import simplejson
        return simplejson.loads(cleaned)
    except:
        pass

    return None

def get_client():
    api_key = st.secrets.get("ANTHROPIC_API_KEY")
    if not api_key:
        st.error("❌ ANTHROPIC_API_KEY가 secrets에 없습니다.")
        return None
    return anthropic.Anthropic(api_key=api_key)

def call_claude(client, prompt, max_tokens=6000, system=None, retries=2):
    """Claude API 호출. max_tokens 잘림 시 토큰을 50% 증량하여 자동 재시도."""
    current_tokens = max_tokens
    for attempt in range(1 + retries):
        try:
            kwargs = dict(
                model="claude-opus-4-6",
                max_tokens=current_tokens,
                messages=[{"role": "user", "content": prompt}]
            )
            if system:
                kwargs["system"] = system
            message = client.messages.create(**kwargs)

            if message.stop_reason == "max_tokens":
                if attempt < retries:
                    next_tokens = min(int(current_tokens * 1.5), 32000)
                    st.info(f"🔄 응답이 {current_tokens} 토큰에서 잘렸습니다. {next_tokens} 토큰으로 재시도합니다... ({attempt+2}/{1+retries}회차)")
                    current_tokens = next_tokens
                    continue
                else:
                    st.warning(f"⚠️ {retries}회 재시도 후에도 응답이 {current_tokens} 토큰에서 잘렸습니다. 결과가 불완전할 수 있습니다.")
            return message.content[0].text
        except Exception as e:
            st.error(f"API 오류 상세: {type(e).__name__} — {e}")
            return None
    return None

# =================================================================
# [3] PDF 추출
# =================================================================
def extract_text(uploaded_file):
    try:
        reader = PdfReader(uploaded_file)
        pages = [p.extract_text() for p in reader.pages if p.extract_text()]
        text = "\n".join(pages)
        if len(text) < 300:
            st.error("텍스트가 너무 짧습니다. 이미지 스캔 PDF는 분석 불가합니다.")
            return None
        return text[:80000]
    except Exception as e:
        st.error(f"PDF 추출 실패: {e}")
        return None

# =================================================================
# [4] BLUE — 분석 비서 (prompt.py 빌더 사용)
# =================================================================
def run_blue(text, client):
    from prompt import SYSTEM_PROMPT
    prompt = build_analysis_prompt(text)
    # SYSTEM_PROMPT가 user 메시지 안에 포함되어 있으므로 제거 후 system으로 분리
    user_prompt = prompt.replace(SYSTEM_PROMPT, '').strip()
    raw = call_claude(client, user_prompt, max_tokens=12000, system=SYSTEM_PROMPT)
    if not raw:
        st.error("❌ API 응답이 없습니다.")
        return None
    data = parse_json(raw)
    if not data:
        st.error(f"❌ JSON 파싱 실패. 응답 앞 300자:\n{raw[:300]}")
        return None
    s = data.get('scores', {})
    data['mark'] = {'final': round(
        s.get('structure', 0) * 0.3 + s.get('hero', 0) * 0.3 +
        s.get('concept', 0) * 0.2 + s.get('genre', 0) * 0.2, 1
    )}
    return data

def _run_blue_OLD_UNUSED(text, client):
    prompt = f"""
당신은 글로벌 스튜디오의 수석 각본 분석가(Senior Script Consultant)입니다.
전 세계 시청자를 사로잡을 수 있는 글로벌 스탠다드 시각으로 시나리오를 정밀 진단하세요.
아래 시나리오를 분석하여 JSON만 출력하세요. 마크다운 금지.
대사 안 쌍따옴표 금지 → 홑따옴표 사용.
tension_data는 숫자만.

[JSON Schema]
{{
  "title": "본문에서 추출한 작품 제목",
  "scores": {{"structure":6,"hero":5,"concept":7,"genre":5}},
  // 각 항목은 반드시 0~10 사이 정수로만 입력. 절대로 100점 만점 사용 금지. 소수점 금지.
  "verdict": {{"status":"추천/고려/비추천","rationale":"3줄 근거"}},
  "logline": {{"original":"원본 로그라인","washed":"개선 로그라인"}},
  "synopsis": "기승전결 줄거리 5~7문장",
  "pros_cons": {{
    "pros":["장점1","장점2","장점3"],
    "cons":["보완점1","보완점2","보완점3"],
    "prescription":"장점을 살리고 보완점을 해결하기 위한 핵심처방 3줄. 구체적 액션 중심으로."
  }},
  "drive": {{
    "goal":"주인공의 목적(욕망)",
    "lack":"발생요인(상실/결핍)",
    "strategy":"해결전략(방법)",
    "evaluation":{{
      "clarity":"목적(욕망)의 명확성 평가 (1-10점 및 이유)",
      "urgency":"발생요인의 확실성 평가 (1-10점 및 이유)",
      "consistency":"해결방법의 창의성과 개연성 평가 (1-10점 및 이유)",
      "overall_diagnosis":"서사 동력 총평"
    }}
  }},
  "beats": {{
    "01. Opening Image":"내용","02. Theme Stated":"내용","03. Set-up":"내용",
    "04. Catalyst":"내용","05. Debate":"내용","06. Break Into Two":"내용",
    "07. B Story":"내용","08. Fun and Games":"내용","09. Midpoint":"내용",
    "10. Bad Guys Close In":"내용","11. All Is Lost":"내용",
    "12. Dark Night of the Soul":"내용","13. Break Into Three":"내용",
    "14. Finale":"내용","15. Final Image":"내용"
  }},
  "characters": {{"names":["인물1","인물2"],"ratios":[60,40]}},
  "tension_data": [0,2,5,8,10,8,5,7,9,10,9,8],
  "genre_suitability": {{
    "genre_name": "장르명(예: 스릴러, 로맨스, 액션 등)",
    "checks": ["장르 필수 문법 1 - 충족 여부", "장르 필수 문법 2 - 충족 여부", "장르 필수 문법 3 - 충족 여부"],
    "compliance_score": 0,
    "missing_elements": ["이 장르에서 반드시 있어야 하는데 빠진 요소1", "빠진 요소2"],
    "doctoring": "이 작품이 해당 장르의 문법을 얼마나 지키고 있는지, 어떤 장르 관습을 위반했는지, 장르 팬이 실망할 포인트가 무엇인지 구체적으로 서술. 스토리 내용 요약 금지."
  }}
}}

시나리오:
{text[:60000]}
"""
    raw = call_claude(client, prompt)
    data = parse_json(raw)
    if data:
        s = data.get('scores', {})
        data['mark'] = {'final': round(
            s.get('structure', 0) * 0.3 + s.get('hero', 0) * 0.3 +
            s.get('concept', 0) * 0.2 + s.get('genre', 0) * 0.2, 1
        )}
    return data

# =================================================================
# [5] JEAN — 워싱 비서 (prompt.py 빌더 사용)
# =================================================================
def run_jean(text, analysis, client):
    from prompt import SYSTEM_PROMPT
    prompt = build_doctoring_prompt(text, analysis)
    user_prompt = prompt.replace(SYSTEM_PROMPT, '').strip()
    raw = call_claude(client, user_prompt, max_tokens=12000, system=SYSTEM_PROMPT)
    return parse_json(raw)

def _run_jean_OLD_UNUSED(text, analysis, client):
    prompt = f"""당신은 글로벌 OTT와 극장 영화 양쪽에서 검증된 세계 최고의 쇼러너(Showrunner)입니다.
작품: {analysis.get('title', '')}
장르: {analysis.get('genre_compliance', {}).get('genre_key', analysis.get('genre_suitability', {}).get('genre_name', ''))}

시나리오의 시퀀스를 진단하고, 대사를 4축으로 심층 분석하고, 각색 제안을 작성하세요.
JSON만 출력. 마크다운 금지.

[washing_table 규칙]
- 라벨은 "동사 → 동사" 형식 필수 (예: "잠복 → 발각")
- 최소 6개 이상
- 처방 끝에 반드시 [Risk: 수정 시 우려점] 기재

[Dialogue Washing — 4축 분석 규칙]
반드시 아래 4가지 기준으로 대사를 분석한다. 각 축마다 실제 대사를 인용하고 개선안을 제시한다.

① 캐릭터 적합성 (Character Voice)
   - 이 인물이 쓸 법한 어휘·문체·말투인가?
   - 인물의 직업·나이·교육수준·감정 상태가 대사에 반영되어 있는가?
   - 문제: 캐릭터 구분 없이 작가 목소리로 통일된 대사

② 서브텍스트 (Subtext)
   - 대사 표면과 이면의 감정이 충돌하는가?
   - 말하지 않는 것이 말하는 것보다 강한가?
   - 문제: 감정을 직접 설명하는 대사 ("나 너 좋아해", "화났어" 같은 설명형)

③ 행동/감정/관계 구동 (Action-Driven)
   - 대사가 장면을 앞으로 밀고 나가는가?
   - 설명·정보 전달이 아닌 행동·감정·관계 변화를 만드는가?
   - 문제: 플롯 설명 대사, 방백형 독백, 과거 회상 나레이션

④ 개선 제안 (Rewrite Suggestion)
   - 위 3축 문제를 동시에 해결하는 구체적인 대사 개선안 제시
   - 세계 정상급 각본 스타일: 빠른 리듬, 반박, 말 끊기, 의미의 역전
   - 홑따옴표 사용 필수. 쌍따옴표 금지.

각 issues 항목의 type은 반드시 다음 중 하나:
"캐릭터 부적합" / "서브텍스트 부재" / "설명형 대사" / "관계 미구동"

[JSON Schema]
{{
  "washing_table": [
    {{"seq":"Seq 1","label":"잠복 → 발각","diagnosis":"진단 2문장 이내","prescription":"처방. [Risk: 우려점]"}}
  ],
  "dialogue_analysis": {{
    "overall_score": 6,
    "overall_comment": "전체 대사 수준 총평. 장르/캐릭터 적합성 포함. 3줄.",
    "axis_scores": {{
      "character_voice": 5,
      "subtext": 6,
      "action_driven": 7
    }},
    "issues": [
      {{
        "type": "서브텍스트 부재",
        "axis": "② 서브텍스트",
        "description": "문제 설명 1~2문장",
        "example_bad": "인물명: '문제 대사'",
        "example_good": "인물명: '개선된 대사'",
        "rewrite_note": "리라이팅 지시 1줄"
      }}
    ],
    "strengths": ["강점 1", "강점 2"],
    "rewrite_directives": [
      "캐릭터 보이스 관련 지시",
      "서브텍스트 관련 지시",
      "행동 구동 관련 지시"
    ]
  }},
  "suggestions": ["1. 제안","2. 제안","3. 제안","4. 제안","5. 제안",
                  "6. 제안","7. 제안","8. 제안","9. 제안","10. 제안"]
}}

시나리오:
{text[:40000]}
"""
    raw = call_claude(client, prompt, max_tokens=8000)
    return parse_json(raw)

# =================================================================
# [6] PICTURES — 리라이팅 비서 (prompt.py 빌더 사용)
# =================================================================
def run_pictures(text, washing, client):
    from prompt import SYSTEM_PROMPT
    analysis = st.session_state.get('analysis', {})
    prompt = build_rewrite_prompt(text, analysis, washing)
    user_prompt = prompt.replace(SYSTEM_PROMPT, '').strip()
    raw = call_claude(client, user_prompt, max_tokens=16000, system=SYSTEM_PROMPT)
    result = parse_json(raw)
    if result and not result.get('rewriting', {}).get('scenes'):
        st.warning(f"⚠️ scenes 파싱 실패. raw 앞 500자: {raw[:500] if raw else 'None'}")
    return result

def _run_pictures_OLD_UNUSED(text, washing, client):
    suggestions = washing.get('suggestions', [])
    washing_table = washing.get('washing_table', [])
    dialogue = washing.get('dialogue_analysis', {})
    
    # 각색 제안 요약
    suggestions_text = "\n".join([f"- {s}" for s in suggestions[:10]])
    
    # 워싱 테이블에서 핵심 문제 추출
    problems_text = "\n".join([
        f"- [{r.get('seq','')}] {r.get('label','')}: {r.get('diagnosis','')}"
        for r in washing_table[:6]
    ])

    # Shiho가 전달한 대사 리라이팅 지시사항
    dialogue_directives = dialogue.get('rewrite_directives', [])
    dialogue_issues = dialogue.get('issues', [])
    
    directives_text = "\n".join([f"- {d}" for d in dialogue_directives])
    issues_text = "\n".join([
        f"- {issue.get('type','')}: {issue.get('rewrite_note','')}\n"
        f"  ❌ {issue.get('example_bad','')}\n"
        f"  ✅ 방향: {issue.get('example_good','')}"
        for issue in dialogue_issues[:4]
    ])

    # 예시 씬 내용 (f-string 밖에서 정의)
    example_revised = (
        "INT. 주인공 집 거실 - 새벽 3시\n\n"
        "텅 빈 소파. 식어버린 커피잔.\n\n"
        "주인공: (전화기를 바라보며) '전화하면 안 되는 사람한테 전화하고 싶은 밤이 있잖아.'\n\n"
        "침묵. 그는 전화기를 뒤집어 화면이 보이지 않게 둔다."
    )
    example_added = (
        "EXT. 한강 다리 위 - 저녁\n\n"
        "바람이 세다. 두 사람이 난간에 기대 서울 야경을 바라본다.\n\n"
        "상대: '왜 그때 말 안 했어?'\n\n"
        "주인공: (강물을 내려다보며) '말했으면 달라졌을 것 같아서. 근데 달라지는 게 더 무서웠거든.'\n\n"
        "상대가 아무 말 없이 손을 잡는다. 주인공은 뿌리치지 않는다."
    )
    fallback_directives = "- 서브텍스트를 담은 대사 작성\n- 인물 고유의 말투 살리기\n- 갈등이 대사 안에 내재될 것"
    fallback_issues = "- 전반적인 대사 품질 향상 필요"

    prompt = f"""당신은 세계 최고의 각본가입니다.
글로벌 스튜디오와 OTT 오리지널 양쪽에서 검증된 각본 리라이터입니다.

아래 시나리오의 핵심 문제를 해결하는 각색 원고를 작성하세요.

[진단된 핵심 문제 — Shiho 진단]
{problems_text}

[각색 제안]
{suggestions_text}

[대사 리라이팅 지시사항 — Shiho가 Moon에게 전달]
{directives_text if directives_text else fallback_directives}

[수정이 필요한 대사 유형]
{issues_text if issues_text else fallback_issues}

[리라이팅 규칙 - 반드시 준수]
1. 총 10개 씬을 작성한다. 9개나 11개는 실패.
2. 수정씬 6개: 원본에서 문제가 있는 장면을 완전히 새로 쓴다.
   - 원본의 핵심 정보(인물, 장소, 사건)는 유지
   - 대사와 지문은 100% 새로 쓴다
   - 진단된 문제를 구체적으로 해결해야 한다
3. 추가씬 4개: 시나리오에 없지만 반드시 필요한 장면을 창작한다.
   - 각색 제안에서 언급된 빠진 요소를 채운다
   - 캐릭터 심리, 관계 변화, 복선 등을 보완한다
4. 각 씬 형식:
   - 씬헤딩: INT./EXT. 장소 - 시간대
   - 지문: 3~5줄, 인물의 행동과 감정 묘사
   - 대사: 인물 이름 뒤 콜론, 홑따옴표 안에 대사
   - 각 씬 200~300자 분량
5. 대사 품질 (Shiho 지시사항 반드시 반영):
   - 인물 고유의 말투와 어휘를 살린다
   - 대사 안에 서브텍스트(말 뒤에 숨은 감정)를 담는다
   - "그래", "알아", "맞아" 같은 밋밋한 대사 금지
   - 한 대사 안에 두 인물의 욕망이 충돌하게 쓴다
6. JSON만 출력. 마크다운 코드블록 금지.

[JSON 출력 형식]
{{
  "rewriting": {{
    "target_reason": "이 10개 씬을 선택한 전략적 이유 — 어떤 문제를 어떻게 해결하는지 구체적으로 3~4줄",
    "scenes": [
      {{
        "scene_no": "S#3",
        "type": "수정씬",
        "insert_between": "",
        "original": "기존: 주인공이 집에서 혼자 고민하는 장면. 독백 위주로 내면을 직접 설명함.",
        "content": "{example_revised}"
      }},
      {{
        "scene_no": "S#신규-감정씬",
        "type": "추가씬",
        "insert_between": "S#7과 S#8 사이 (복선 구간)",
        "original": "",
        "content": "{example_added}"
      }}
    ]
  }}
}}

[분석할 시나리오]
{text[:45000]}
"""
    raw = call_claude(client, prompt, max_tokens=16000)
    result = parse_json(raw)
    # scenes가 비어있으면 raw 확인용 (개발 디버그)
    if result and not result.get('rewriting', {}).get('scenes'):
        st.warning(f"⚠️ scenes 파싱 실패. raw 앞 500자: {raw[:500] if raw else 'None'}")
    return result

# =================================================================
# [7] 리포트 렌더러
# =================================================================
def render_analysis(data):

    # 1. 종합
    # ── 4축 점수 계산 ──
    sc       = data.get('scores', {})
    s_score  = sc.get('structure', 0)
    h_score  = sc.get('hero', 0)
    c_score  = sc.get('concept', 0)
    g_score  = sc.get('genre', 0)
    final    = data.get('mark', {}).get('final', 0)
    verdict  = data.get('verdict', {})

    def score_bar(val, color='#191970'):
        pct = min(int(val) * 10, 100)
        return (
            f'<div style="display:flex;align-items:center;gap:8px;">'
            f'<div style="flex:1;background:#E6E9EF;border-radius:20px;height:8px;overflow:hidden;">'
            f'<div style="background:{color};width:{pct}%;height:100%;border-radius:20px;"></div></div>'
            f'<div style="font-size:0.82rem;font-weight:900;color:{color};min-width:28px;">{val}</div>'
            f'</div>'
        )

    verdict_color = {'추천': '#2EC484', '고려': '#FFCB05', '비추천': '#FF6432'}.get(verdict.get('status', ''), '#191970')

    # score_bar 변수 미리 생성 (f-string 안 함수호출 → Streamlit 1.54 이스케이프 버그 방지)
    bar_s = score_bar(s_score, '#191970')
    bar_h = score_bar(h_score, '#191970')
    bar_c = score_bar(c_score, '#B8860B')
    bar_g = score_bar(g_score, '#2EC484')
    final_pct = min(int(float(final)*10), 100)
    verdict_status    = safe(verdict.get('status',''))
    verdict_rationale = safe(verdict.get('rationale',''))

    # 섹션1-A: 점수 헤더 카드
    st.markdown(f"""
    <div class="report-card">
        <h3>1. 종합 분석 (Total Analysis) — Hollywood Standard</h3>
        <div style="display:flex;align-items:center;gap:20px;margin-bottom:20px;
                    background:#F0F2FF;border-radius:12px;padding:16px 20px;">
            <div style="text-align:center;min-width:70px;">
                <div style="font-size:3rem;font-weight:950;color:#FFCB05;line-height:1;">{final}</div>
                <div style="font-size:0.72rem;color:#191970;opacity:0.5;margin-top:2px;">/ 10.0</div>
            </div>
            <div style="width:1px;height:52px;background:#E0E4F0;"></div>
            <div style="flex:1;">
                <div style="display:inline-block;background:{verdict_color};color:#FFFFFF;
                            padding:3px 14px;border-radius:20px;font-weight:900;
                            font-size:0.82rem;margin-bottom:6px;">{verdict_status}</div>
                <div style="font-size:0.86rem;line-height:1.7;color:#191970;">{verdict_rationale}</div>
            </div>
        </div>
        <div style="font-size:0.72rem;font-weight:900;letter-spacing:0.1em;color:#191970;
                    opacity:0.5;margin-bottom:10px;">4축 정밀 평가 (Hollywood Standard)</div>
    </div>""", unsafe_allow_html=True)

    # 섹션1-B: 4축 테이블 (score_bar 변수 사용)
    st.markdown(f"""
    <div class="report-card" style="margin-top:4px;">
        <table style="width:100%;border-collapse:collapse;font-size:0.85rem;">
            <thead>
                <tr style="background:#191970;">
                    <th style="padding:8px 12px;text-align:left;color:#FFCB05 !important;font-size:0.72rem;letter-spacing:0.08em;width:22%;">AXIS</th>
                    <th style="padding:8px 12px;text-align:left;color:#FFCB05 !important;font-size:0.72rem;letter-spacing:0.08em;width:18%;">가중치</th>
                    <th style="padding:8px 12px;text-align:left;color:#FFCB05 !important;font-size:0.72rem;letter-spacing:0.08em;">평가 기준</th>
                    <th style="padding:8px 12px;text-align:left;color:#FFCB05 !important;font-size:0.72rem;letter-spacing:0.08em;width:30%;">점수</th>
                </tr>
            </thead>
            <tbody>
                <tr style="background:#F8F9FF;border-bottom:1px solid #E6E9EF;">
                    <td style="padding:10px 12px;font-weight:900;color:#191970;">① STRUCTURE<br><span style="font-weight:400;font-size:0.75rem;color:#8899BB;">구성/플롯</span></td>
                    <td style="padding:10px 12px;font-weight:700;color:#191970;">30%</td>
                    <td style="padding:10px 12px;color:#191970;font-size:0.82rem;line-height:1.5;">인과관계의 정밀도<br>3막 구조의 완성도</td>
                    <td style="padding:10px 12px;">{bar_s}</td>
                </tr>
                <tr style="background:#FFFFFF;border-bottom:1px solid #E6E9EF;">
                    <td style="padding:10px 12px;font-weight:900;color:#191970;">② HERO<br><span style="font-weight:400;font-size:0.75rem;color:#8899BB;">캐릭터</span></td>
                    <td style="padding:10px 12px;font-weight:700;color:#191970;">30%</td>
                    <td style="padding:10px 12px;color:#191970;font-size:0.82rem;line-height:1.5;">Goal / Need / Strategy<br>감정선의 선명도</td>
                    <td style="padding:10px 12px;">{bar_h}</td>
                </tr>
                <tr style="background:#F8F9FF;border-bottom:1px solid #E6E9EF;">
                    <td style="padding:10px 12px;font-weight:900;color:#191970;">③ CONCEPT<br><span style="font-weight:400;font-size:0.75rem;color:#8899BB;">소재/컨셉</span></td>
                    <td style="padding:10px 12px;font-weight:700;color:#191970;">20%</td>
                    <td style="padding:10px 12px;color:#191970;font-size:0.82rem;line-height:1.5;">하이컨셉 여부 · 독창성<br>시장성 있는 설정</td>
                    <td style="padding:10px 12px;">{bar_c}</td>
                </tr>
                <tr style="background:#FFFFFF;border-bottom:1px solid #E6E9EF;">
                    <td style="padding:10px 12px;font-weight:900;color:#191970;">④ GENRE<br><span style="font-weight:400;font-size:0.75rem;color:#8899BB;">장르 적합성</span></td>
                    <td style="padding:10px 12px;font-weight:700;color:#191970;">20%</td>
                    <td style="padding:10px 12px;color:#191970;font-size:0.82rem;line-height:1.5;">장르 문법 충실도<br>타깃 관객 소구력</td>
                    <td style="padding:10px 12px;">{bar_g}</td>
                </tr>
                <tr style="background:#EEF0FA;">
                    <td colspan="2" style="padding:10px 12px;font-weight:900;color:#191970;font-size:0.9rem;">FINAL</td>
                    <td style="padding:10px 12px;font-size:0.78rem;color:#191970;opacity:0.6;">0.3S + 0.3H + 0.2C + 0.2G</td>
                    <td style="padding:10px 12px;">
                        <div style="display:flex;align-items:center;gap:8px;">
                            <div style="flex:1;background:#C5CBE8;border-radius:20px;height:8px;overflow:hidden;">
                                <div style="background:#FFCB05;width:{final_pct}%;height:100%;border-radius:20px;"></div>
                            </div>
                            <div style="font-size:0.95rem;font-weight:950;color:#FFCB05;min-width:36px;">{final}</div>
                        </div>
                    </td>
                </tr>
            </tbody>
        </table>
    </div>""", unsafe_allow_html=True)

    # 2. 로그라인
    log = data.get('logline', {})
    st.markdown(f"""
    <div class="report-card">
        <h3>2. 로그라인 분석 (Logline Pack)</h3>
        <div style="padding:14px;border:1px solid #E6E9EF;border-radius:8px;margin-bottom:12px;">
            <div style="font-size:0.7rem;font-weight:700;letter-spacing:0.12em;margin-bottom:6px;color:#191970;opacity:0.6;">ORIGINAL</div>
            <div style="line-height:1.7;color:#191970;">{safe(log.get('original', ''))}</div>
        </div>
        <div style="padding:14px;border:2px solid #FFCB05;border-radius:8px;background:#FFFBE6;">
            <div style="font-size:0.7rem;font-weight:700;letter-spacing:0.12em;margin-bottom:6px;color:#B8860B;">✨ WASHED</div>
            <div style="font-weight:700;font-size:1.05rem;line-height:1.7;color:#191970;">{safe(log.get('washed', ''))}</div>
        </div>
    </div>""", unsafe_allow_html=True)

    # 3. 줄거리
    st.markdown(f"""
    <div class="report-card">
        <h3>3. 줄거리 분석 (Synopsis)</h3>
        <div style="line-height:1.9;color:#191970;">{safe(data.get('synopsis', ''))}</div>
    </div>""", unsafe_allow_html=True)

    # 4. 장단점
    pc = data.get('pros_cons', {})
    pros_html = "".join([f"<li style='margin-bottom:7px;line-height:1.6;color:#191970;'>{safe(p)}</li>" for p in pc.get('pros', [])])
    cons_html = "".join([f"<li style='margin-bottom:7px;line-height:1.6;color:#191970;'>{safe(c)}</li>" for c in pc.get('cons', [])])
    prescription = safe(pc.get('prescription', ''))
    st.markdown(f"""
    <div class="report-card">
        <h3>4. 장점 및 보완점 (Pros &amp; Cons)</h3>
        <div style="display:flex;gap:14px;">
            <div style="flex:1;background:#EDFAF3;border-left:3px solid #2EC484;padding:16px;border-radius:8px;">
                <div style="color:#1A7A50;font-weight:800;font-size:0.75rem;margin-bottom:10px;">✅ PROS</div>
                <ul style="margin:0;padding-left:16px;">{pros_html}</ul>
            </div>
            <div style="flex:1;background:#FFF3EE;border-left:3px solid #FF6432;padding:16px;border-radius:8px;">
                <div style="color:#CC3300;font-weight:800;font-size:0.75rem;margin-bottom:10px;">⚠️ CONS</div>
                <ul style="margin:0;padding-left:16px;">{cons_html}</ul>
            </div>
        </div>
    </div>""", unsafe_allow_html=True)
    if prescription:
        st.markdown(f"""
        <div style="background:#FFFBE6;border-left:4px solid #FFCB05;padding:14px 16px;
                    border-radius:8px;margin-top:-8px;margin-bottom:16px;">
            <div style="font-size:0.75rem;font-weight:800;color:#B8860B;
                        margin-bottom:6px;letter-spacing:0.08em;">💊 핵심 처방 (Key Prescription)</div>
            <div style="font-size:0.9rem;line-height:1.8;color:#191970;font-weight:600;">{prescription}</div>
        </div>""", unsafe_allow_html=True)

    # 5. 서사동력
    drive = data.get('drive', {})
    ev = drive.get('evaluation', {})
    st.markdown(f"""
    <div class="report-card">
        <h3>5. 서사 동력 (Narrative Drive)</h3>
        <div style="display:flex;gap:10px;margin-bottom:14px;">
            <div style="flex:1;background:#FFFBE6;padding:14px;border-radius:10px;text-align:center;border:1px solid #FFE066;">
                <div style="font-size:0.72rem;color:#B8860B;font-weight:700;margin-bottom:7px;">① 목적(욕망)</div>
                <div style="font-size:0.87rem;line-height:1.5;color:#191970;">{safe(drive.get('goal', ''))}</div>
            </div>
            <div style="flex:1;background:#FFFBE6;padding:14px;border-radius:10px;text-align:center;border:1px solid #FFE066;">
                <div style="font-size:0.72rem;color:#B8860B;font-weight:700;margin-bottom:7px;">② 발생요인</div>
                <div style="font-size:0.87rem;line-height:1.5;color:#191970;">{safe(drive.get('lack', ''))}</div>
            </div>
            <div style="flex:1;background:#FFFBE6;padding:14px;border-radius:10px;text-align:center;border:1px solid #FFE066;">
                <div style="font-size:0.72rem;color:#B8860B;font-weight:700;margin-bottom:7px;">③ 해결전략</div>
                <div style="font-size:0.87rem;line-height:1.5;color:#191970;">{safe(drive.get('strategy', ''))}</div>
            </div>
        </div>
        <div style="background:#EEF0FA;padding:14px;border-radius:8px;border:1px solid #C5CBE8;">
            <p style="margin:4px 0;font-size:0.88rem;color:#191970;"><strong>목적(욕망) 명확성</strong>  {safe(ev.get('clarity', ''))}</p>
            <p style="margin:4px 0;font-size:0.88rem;color:#191970;"><strong>발생요인 확실성</strong>  {safe(ev.get('urgency', ''))}</p>
            <p style="margin:4px 0;font-size:0.88rem;color:#191970;"><strong>해결전략 창의성</strong>  {safe(ev.get('consistency', ''))}</p>
            <p style="margin:10px 0 0;font-size:0.88rem;line-height:1.7;color:#191970;border-top:1px solid #C5CBE8;padding-top:10px;">{safe(ev.get('overall_diagnosis', ''))}</p>
        </div>
    </div>""", unsafe_allow_html=True)

    # 6. 15-Beat Sheet
    beats = data.get('beats', {})
    circles = {i: c for i, c in enumerate(['①','②','③','④','⑤','⑥','⑦','⑧','⑨','⑩','⑪','⑫','⑬','⑭','⑮'], 1)}
    BEAT_KO = {
        "Opening Image":   "오프닝 이미지",
        "Theme Stated":    "주제 제시",
        "Set-up":          "설정",
        "Catalyst":        "촉매",
        "Debate":          "갈등",
        "Break Into Two":  "2막 진입",
        "B Story":         "B 스토리",
        "Fun and Games":   "재미와 게임",
        "Midpoint":        "중간점",
        "Bad Guys Close In":"위기 고조",
        "All Is Lost":     "모든 것을 잃다",
        "Dark Night of Soul":"영혼의 어둔 밤",
        "Break Into Three":"3막 진입",
        "Finale":          "피날레",
        "Final Image":     "최종 이미지",
    }
    rows = ""
    for idx, (k, v) in enumerate(sorted(beats.items()), 1):
        name = re.sub(r'^[\d\.\-\_\s]+', '', str(k)).strip()
        ko = BEAT_KO.get(name, '')
        display_name = f"{ko}<br><span style='font-weight:400;font-size:0.75rem;color:#8899BB;'>({name})</span>" if ko else name
        bg = "#F8F9FF" if idx % 2 == 0 else "#FFFFFF"
        rows += (f'<tr style="background:{bg};border-bottom:1px solid #E6E9EF;">'
                 f'<td style="padding:9px 14px;font-weight:700;color:#191970;width:26%;font-size:0.85rem;">{circles.get(idx,"")} {display_name}</td>'
                 f'<td style="padding:9px 14px;line-height:1.6;font-size:0.86rem;color:#191970;">{safe(v)}</td></tr>')
    st.markdown(f"""
    <div class="report-card">
        <h3>6. 구성 및 플롯 (15-Beat Sheet)</h3>
        <table style="width:100%;border-collapse:collapse;">
            <thead><tr style="background:#191970;">
                <th style="padding:9px 14px;text-align:left;font-size:0.75rem;letter-spacing:0.1em;color:#FFCB05 !important;width:26%;">BEAT</th>
                <th style="padding:9px 14px;text-align:left;font-size:0.75rem;letter-spacing:0.1em;color:#FFCB05 !important;">DESCRIPTION</th>
            </tr></thead>
            <tbody>{rows}</tbody>
        </table>
    </div>""", unsafe_allow_html=True)

    # 7. 시각화
    st.markdown('<div class="report-card"><h3>7. 시각화 (Visualization)</h3></div>', unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    with col1:
        arc = data.get('tension_data', [])
        nums = []
        for v in arc:
            try: nums.append(float(v))
            except: pass
        if nums:
            fig = px.line(y=nums, template="plotly_white", title="긴장도 곡선 (Tension Arc)")
            fig.update_traces(line_color='#191970', line_width=2.5, mode='lines+markers',
                              marker=dict(size=6, color='#FFCB05', line=dict(color='#191970', width=1)))
            fig.update_layout(height=260, margin=dict(l=10,r=10,t=36,b=10),
                              paper_bgcolor='#FAFAFA', plot_bgcolor='#FAFAFA',
                              yaxis=dict(range=[0,11], gridcolor='#E6E9EF'),
                              xaxis=dict(gridcolor='#E6E9EF'),
                              title_font_color='#191970', font_color='#191970', title_font_size=13)
            st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': False})
    with col2:
        chars = data.get('characters', {})
        names = chars.get('names', [])
        ratios = []
        for r in chars.get('ratios', []):
            try: ratios.append(float(r))
            except: ratios.append(0)
        if names and ratios:
            fig2 = px.pie(values=ratios[:len(names)], names=names[:len(ratios)],
                          hole=0.45, title="인물 비중 (Narrative Share)",
                          color_discrete_sequence=['#191970','#FFCB05','#2EC484','#FF6432','#94A3B8'])
            fig2.update_layout(height=260, margin=dict(l=10,r=10,t=36,b=10),
                               paper_bgcolor='#FAFAFA',
                               title_font_color='#191970', font_color='#191970', title_font_size=13,
                               legend=dict(font=dict(color='#191970', size=11)))
            st.plotly_chart(fig2, use_container_width=True, config={'displayModeBar': False})

    # 8. 장르
    genre = data.get('genre_compliance', {})
    # must_have_check 렌더링 (새 스키마)
    must_checks = genre.get('must_have_check', [])
    if must_checks:
        tags = "".join([
            f"<span style='background:{'#EDFAF3' if c.get('status','') == '충족' else ('#FFFBE6' if c.get('status','') == '약함' else '#FFF3EE')};color:{'#1A7A50' if c.get('status','') == '충족' else ('#B8860B' if c.get('status','') == '약함' else '#CC3300')};border:1px solid {'#2EC484' if c.get('status','') == '충족' else ('#FFE066' if c.get('status','') == '약함' else '#FF6432')};padding:3px 11px;border-radius:20px;font-size:0.75rem;font-weight:700;margin-right:7px;margin-bottom:6px;display:inline-block;'>{'✓' if c.get('status','') == '충족' else ('△' if c.get('status','') == '약함' else '✗')} {safe(c.get('item',''))}</span>"
            for c in must_checks
        ])
    else:
        # fallback: 이전 스키마 호환
        tags = "".join([
            f"<span style='background:#EDFAF3;color:#1A7A50;border:1px solid #2EC484;padding:3px 11px;border-radius:20px;font-size:0.75rem;font-weight:700;margin-right:7px;margin-bottom:6px;display:inline-block;'>✓ {safe(c)}</span>"
            for c in genre.get('checks', [])
        ])
    missing = genre.get('missing_elements', [])
    missing_html = "".join([
        f"<span style='background:#FFF3EE;color:#CC3300;border:1px solid #FF6432;padding:3px 10px;border-radius:20px;font-size:0.75rem;font-weight:700;margin-right:6px;margin-bottom:6px;display:inline-block;'>✗ {safe(m)}</span>"
        for m in missing
    ])
    compliance = genre.get('compliance_score', 0)
    bar_color = '#2EC484' if compliance >= 7 else ('#FFCB05' if compliance >= 4 else '#FF6432')
    genre_display = safe(genre.get('genre_key', genre.get('genre_name', '')))
    
    # 장르적 재미 진단
    genre_fun_diagnosis = safe(genre.get('genre_fun_diagnosis', ''))
    genre_fun_alive = genre.get('genre_fun_alive', None)
    fun_badge = ""
    if genre_fun_alive is not None:
        fun_badge = (f"<span style='background:{'#EDFAF3' if genre_fun_alive else '#FFF3EE'};color:{'#1A7A50' if genre_fun_alive else '#CC3300'};"
                     f"border:1px solid {'#2EC484' if genre_fun_alive else '#FF6432'};padding:3px 11px;border-radius:20px;font-size:0.75rem;"
                     f"font-weight:700;margin-left:10px;'>{'✓ 장르적 재미 작동' if genre_fun_alive else '✗ 장르적 재미 약함'}</span>")

    # 실패 패턴 데이터 (렌더링은 아래에서 별도 st.markdown으로)
    fail_patterns = genre.get('fail_patterns_found', [])

    st.markdown(f"""
    <div class="report-card">
        <h3>8. 장르 분석 및 적합도 (Genre Compliance)</h3>
        <div style="display:flex;align-items:center;gap:12px;margin-bottom:14px;">
            <div style="font-size:1.5rem;font-weight:900;color:#191970;">{genre_display}</div>
            {fun_badge}
            <div style="flex:1;background:#E6E9EF;border-radius:20px;height:10px;overflow:hidden;">
                <div style="background:{bar_color};width:{compliance*10}%;height:100%;border-radius:20px;"></div>
            </div>
            <div style="font-size:1.1rem;font-weight:900;color:{bar_color};">{compliance}/10</div>
        </div>
        <div style="margin-bottom:10px;"><strong style="font-size:0.78rem;color:#191970;display:block;margin-bottom:6px;">✅ 장르 필수 요소 체크</strong>{tags}</div>
        {"<div style='margin-bottom:12px;'><strong style='font-size:0.78rem;color:#CC3300;display:block;margin-bottom:6px;'>❌ 누락된 장르 필수 요소</strong>" + missing_html + "</div>" if missing else ""}
    </div>""", unsafe_allow_html=True)

    # 실패 패턴 (별도 렌더링)
    if fail_patterns:
        fail_items = "".join([
            f"<span style='background:#FFF3EE;color:#CC3300;border:1px solid #FF6432;padding:3px 10px;border-radius:20px;font-size:0.75rem;font-weight:700;margin-right:6px;margin-bottom:6px;display:inline-block;'>⚠ {safe(f)}</span>"
            for f in fail_patterns
        ])
        st.markdown(f"""
        <div class="report-card" style="margin-top:-8px;">
            <strong style="font-size:0.78rem;color:#CC3300;display:block;margin-bottom:6px;">⚠️ 발견된 실패 패턴</strong>
            {fail_items}
        </div>""", unsafe_allow_html=True)

    # Hook/Punch 체크 (별도 렌더링)
    hp = genre.get('hook_punch_check', {})
    if hp:
        hook_ok = hp.get('hook_present', False)
        punch_ok = hp.get('punch_present', False)
        hook_bg = '#EDFAF3' if hook_ok else '#FFF3EE'
        hook_border = '#2EC484' if hook_ok else '#FF6432'
        hook_color = '#1A7A50' if hook_ok else '#CC3300'
        hook_icon = '✓' if hook_ok else '✗'
        punch_bg = '#EDFAF3' if punch_ok else '#FFF3EE'
        punch_border = '#2EC484' if punch_ok else '#FF6432'
        punch_color = '#1A7A50' if punch_ok else '#CC3300'
        punch_icon = '✓' if punch_ok else '✗'
        hook_note = safe(hp.get('hook_note', ''))
        punch_note = safe(hp.get('punch_note', ''))
        st.markdown(f"""
        <div class="report-card" style="margin-top:-8px;">
            <div style="display:flex;gap:10px;">
                <div style="flex:1;background:{hook_bg};padding:10px;border-radius:8px;border-left:3px solid {hook_border};">
                    <div style="font-size:0.72rem;font-weight:800;color:{hook_color};margin-bottom:4px;">{hook_icon} Hook (오프닝)</div>
                    <div style="font-size:0.84rem;color:#191970;line-height:1.6;">{hook_note}</div>
                </div>
                <div style="flex:1;background:{punch_bg};padding:10px;border-radius:8px;border-left:3px solid {punch_border};">
                    <div style="font-size:0.72rem;font-weight:800;color:{punch_color};margin-bottom:4px;">{punch_icon} Punch (결정타)</div>
                    <div style="font-size:0.84rem;color:#191970;line-height:1.6;">{punch_note}</div>
                </div>
            </div>
        </div>""", unsafe_allow_html=True)

    # 장르적 재미 진단 (별도 렌더링)
    if genre_fun_diagnosis:
        st.markdown(f"""
        <div class="report-card" style="margin-top:-8px;">
            <div style="background:#FFFBE6;border-left:3px solid #FFCB05;padding:14px;border-radius:8px;line-height:1.8;color:#191970;">
                <strong style="font-size:0.75rem;color:#B8860B;display:block;margin-bottom:6px;">🎬 장르적 재미 진단</strong>
                {genre_fun_diagnosis}
            </div>
        </div>""", unsafe_allow_html=True)

    # 장르 종합 코멘트
    doctoring = safe(genre.get('doctoring', ''))
    if doctoring:
        st.markdown(f"""
        <div class="report-card" style="margin-top:-8px;">
            <div style="background:#EEF0FA;padding:16px;border-radius:8px;line-height:1.8;color:#191970;">{doctoring}</div>
        </div>""", unsafe_allow_html=True)


def render_washing(data):
    # 9. 시퀀스 워싱
    st.markdown('<div class="report-card"><h3>9. 시퀀스 워싱 (Washing Table)</h3></div>', unsafe_allow_html=True)
    for row in data.get('washing_table', []):
        seq          = safe(row.get('seq', ''))
        label        = safe(row.get('label', ''))
        diagnosis    = safe(row.get('diagnosis', ''))
        prescription = safe(row.get('prescription', ''))
        st.markdown(f"""
        <div style="background:#FFFFFF;border:1px solid #E6E9EF;border-radius:10px;padding:16px;margin-bottom:12px;">
            <div style="margin-bottom:10px;">
                <span style="background:#191970;color:#FFFFFF !important;padding:2px 10px;border-radius:4px;font-weight:900;font-size:0.72rem;">{seq}</span>
                <span style="font-weight:700;margin-left:10px;color:#191970;">{label}</span>
            </div>
            <div style="display:flex;gap:12px;">
                <div style="flex:1;background:#FFF5F5;padding:12px;border-radius:8px;border-left:3px solid #D32F2F;">
                    <div style="color:#D32F2F;font-size:0.72rem;font-weight:800;margin-bottom:5px;">⚠️ 진단</div>
                    <div style="font-size:0.88rem;line-height:1.6;color:#191970;">{diagnosis}</div>
                </div>
                <div style="flex:1.2;background:#EEF0FA;padding:12px;border-radius:8px;border-left:3px solid #191970;">
                    <div style="color:#191970;font-size:0.72rem;font-weight:800;margin-bottom:5px;">✅ 처방</div>
                    <div style="font-size:0.88rem;line-height:1.6;color:#191970;">{prescription}</div>
                </div>
            </div>
        </div>""", unsafe_allow_html=True)

    # 9-B. 대사 분석 (Dialogue Washing)
    da = data.get('dialogue_analysis', {})
    if da:
        # 점수 정규화: AI가 0~100으로 줄 수도 있으므로 10 초과면 /10
        def norm(v):
            v = float(v or 0)
            return round(v / 10, 1) if v > 10 else v

        score = norm(da.get('overall_score', 0))
        bar_color = '#2EC484' if score >= 7 else ('#FFCB05' if score >= 4 else '#FF6432')
        ax = da.get('axis_scores', {})
        cv  = norm(ax.get('character_voice', 0))
        st_ = norm(ax.get('subtext', 0))
        ad  = norm(ax.get('action_driven', 0))

        def mini_bar(val, color):
            pct = min(int(float(val) * 10), 100)
            return (f'<div style="display:flex;align-items:center;gap:6px;">'
                    f'<div style="flex:1;background:#E6E9EF;border-radius:10px;height:6px;overflow:hidden;">'
                    f'<div style="background:{color};width:{pct}%;height:100%;border-radius:10px;"></div></div>'
                    f'<span style="font-size:0.78rem;font-weight:900;color:{color};">{val}</span></div>')

        # mini_bar 미리 생성 (f-string 안에서 함수 호출 방지)
        bar_cv = mini_bar(cv, '#191970')
        bar_st = mini_bar(st_, '#6B3FA0')
        bar_ad = mini_bar(ad, '#2EC484')
        overall_comment = safe(da.get('overall_comment',''))
        bar_final = (
            f'<div style="display:flex;align-items:center;gap:6px;">'
            f'<div style="flex:1;background:#C5CBE8;border-radius:10px;height:8px;overflow:hidden;">'
            f'<div style="background:{bar_color};width:{min(int(float(score)*10),100)}%;height:100%;border-radius:10px;"></div></div>'
            f'<span style="font-size:0.9rem;font-weight:950;color:{bar_color};">{score}/10</span></div>'
        )

        # 헤더 + 4축 기준 태그
        st.markdown("""
        <div class="report-card">
            <h3>9-B. 대사 워싱 (Dialogue Washing)</h3>
            <div style="display:flex;gap:6px;flex-wrap:wrap;margin-bottom:14px;">
                <span style="background:#EEF0FA;color:#191970;border:1px solid #C5CBE8;padding:3px 10px;border-radius:20px;font-size:0.72rem;font-weight:700;">① 캐릭터 적합성</span>
                <span style="background:#EEF0FA;color:#191970;border:1px solid #C5CBE8;padding:3px 10px;border-radius:20px;font-size:0.72rem;font-weight:700;">② 서브텍스트</span>
                <span style="background:#EEF0FA;color:#191970;border:1px solid #C5CBE8;padding:3px 10px;border-radius:20px;font-size:0.72rem;font-weight:700;">③ 행동/감정/관계 구동</span>
                <span style="background:#FFFBE6;color:#B8860B;border:1px solid #FFE066;padding:3px 10px;border-radius:20px;font-size:0.72rem;font-weight:700;">④ 개선 제안</span>
            </div>
        </div>""", unsafe_allow_html=True)

        # 3축 점수 테이블 (별도 markdown)
        st.markdown(f"""
        <div style="margin-bottom:12px;">
        <table style="width:100%;border-collapse:collapse;margin-bottom:10px;">
            <thead><tr style="background:#191970;">
                <th style="padding:7px 12px;text-align:left;color:#FFCB05 !important;font-size:0.7rem;letter-spacing:0.08em;">평가 축</th>
                <th style="padding:7px 12px;text-align:left;color:#FFCB05 !important;font-size:0.7rem;letter-spacing:0.08em;">기준</th>
                <th style="padding:7px 12px;text-align:left;color:#FFCB05 !important;font-size:0.7rem;letter-spacing:0.08em;width:35%;">점수</th>
            </tr></thead>
            <tbody>
                <tr style="background:#F8F9FF;border-bottom:1px solid #E6E9EF;">
                    <td style="padding:8px 12px;font-weight:800;color:#191970;font-size:0.82rem;">① 캐릭터 적합성</td>
                    <td style="padding:8px 12px;font-size:0.78rem;color:#191970;">고유 어휘·말투·감정 반영</td>
                    <td style="padding:8px 12px;">{bar_cv}</td>
                </tr>
                <tr style="background:#FFFFFF;border-bottom:1px solid #E6E9EF;">
                    <td style="padding:8px 12px;font-weight:800;color:#191970;font-size:0.82rem;">② 서브텍스트</td>
                    <td style="padding:8px 12px;font-size:0.78rem;color:#191970;">표면↔이면 충돌, 설명형 금지</td>
                    <td style="padding:8px 12px;">{bar_st}</td>
                </tr>
                <tr style="background:#F8F9FF;border-bottom:1px solid #E6E9EF;">
                    <td style="padding:8px 12px;font-weight:800;color:#191970;font-size:0.82rem;">③ 행동/감정/관계</td>
                    <td style="padding:8px 12px;font-size:0.78rem;color:#191970;">장면 추진력, 정보전달 금지</td>
                    <td style="padding:8px 12px;">{bar_ad}</td>
                </tr>
                <tr style="background:#EEF0FA;">
                    <td colspan="2" style="padding:8px 12px;font-weight:900;color:#191970;font-size:0.88rem;">종합 대사 수준</td>
                    <td style="padding:8px 12px;">{bar_final}</td>
                </tr>
            </tbody>
        </table>
        <div style="background:#EEF0FA;padding:14px;border-radius:8px;line-height:1.8;color:#191970;">
            {overall_comment}
        </div>
        </div>""", unsafe_allow_html=True)

        # 강점
        strengths = da.get('strengths', [])
        if strengths:
            s_html = "".join([
                f"<span style='background:#EDFAF3;color:#1A7A50;border:1px solid #2EC484;padding:3px 11px;border-radius:20px;font-size:0.75rem;font-weight:700;margin-right:7px;margin-bottom:6px;display:inline-block;'>✓ {safe(s)}</span>"
                for s in strengths
            ])
            st.markdown(f"""
            <div style="background:#FFFFFF;border:1px solid #E6E9EF;border-radius:10px;padding:14px;margin-bottom:12px;">
                <div style="font-size:0.75rem;font-weight:800;color:#1A7A50;margin-bottom:8px;">💪 대사 강점</div>
                <div>{s_html}</div>
            </div>""", unsafe_allow_html=True)

        # 4축 기준별 Before/After
        issues = da.get('issues', [])
        if issues:
            st.markdown("""
            <div style="font-size:0.85rem;font-weight:800;color:#CC3300;margin:16px 0 8px;">
                🔍 Dialogue Washing — 4축 진단 Before / After
            </div>""", unsafe_allow_html=True)
            axis_colors = {
                '① 캐릭터 적합성': '#191970',
                '② 서브텍스트':    '#6B3FA0',
                '③ 행동/감정/관계':'#2EC484',
            }
            for issue in issues:
                axis_label = safe(issue.get('axis', ''))
                axis_color = axis_colors.get(issue.get('axis', ''), '#FF6432')
                st.markdown(f"""
                <div style="background:#FFFFFF;border:1px solid #E6E9EF;border-radius:10px;padding:16px;margin-bottom:12px;">
                    <div style="display:flex;align-items:center;gap:8px;margin-bottom:10px;">
                        <span style="background:#FF6432;color:#FFFFFF;padding:2px 10px;border-radius:4px;font-weight:900;font-size:0.72rem;">{safe(issue.get('type',''))}</span>
                        <span style="background:{axis_color};color:#FFFFFF;padding:2px 10px;border-radius:4px;font-weight:700;font-size:0.68rem;">{axis_label}</span>
                        <span style="font-size:0.83rem;color:#555;">{safe(issue.get('description',''))}</span>
                    </div>
                    <div style="display:flex;gap:10px;margin-bottom:10px;">
                        <div style="flex:1;background:#FFF5F5;padding:11px;border-radius:8px;border-left:3px solid #FF6432;">
                            <div style="font-size:0.68rem;font-weight:800;color:#CC3300;margin-bottom:5px;">❌ BEFORE (원문)</div>
                            <div style="font-size:0.88rem;line-height:1.6;color:#333;font-family:'Courier New',monospace;">{safe(issue.get('example_bad',''))}</div>
                        </div>
                        <div style="flex:1;background:#EDFAF3;padding:11px;border-radius:8px;border-left:3px solid #2EC484;">
                            <div style="font-size:0.68rem;font-weight:800;color:#1A7A50;margin-bottom:5px;">✅ ④ 개선 제안 (AFTER)</div>
                            <div style="font-size:0.88rem;line-height:1.6;color:#333;font-family:'Courier New',monospace;">{safe(issue.get('example_good',''))}</div>
                        </div>
                    </div>
                    <div style="background:#FFFBE6;padding:10px;border-radius:6px;border-left:3px solid #FFCB05;">
                        <div style="font-size:0.68rem;font-weight:800;color:#B8860B;margin-bottom:4px;">✏️ Moon 리라이팅 지시</div>
                        <div style="font-size:0.85rem;color:#191970;">{safe(issue.get('rewrite_note',''))}</div>
                    </div>
                </div>""", unsafe_allow_html=True)

        # rewrite_directives는 내부 프롬프트용 - 보고서에 미노출

    # 10. 각색 제안
    suggestions = data.get('suggestions', [])
    st.markdown('<div class="report-card"><h3>10. 각색 제안 (Action Plan)</h3>', unsafe_allow_html=True)
    for i, s in enumerate(suggestions, 1):
        clean = re.sub(r'^[\d\.\s]+', '', str(s)).strip()
        st.markdown(f"""
        <div style="background:#F8F9FF;border:1px solid #E6E9EF;border-radius:8px;padding:12px;display:flex;gap:10px;align-items:flex-start;margin-bottom:8px;">
            <div style="background:#191970;color:#FFFFFF !important;border-radius:5px;padding:2px 8px;font-weight:900;font-size:0.68rem;white-space:nowrap;flex-shrink:0;margin-top:2px;">STEP {i:02d}</div>
            <div style="font-size:0.88rem;line-height:1.5;color:#191970;">{safe(clean)}</div>
        </div>""", unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)


def render_rewriting(data):
    rewrite = data.get('rewriting', {})
    scenes = rewrite.get('scenes', [])
    revised_count = sum(1 for s in scenes if s.get('type') == '수정씬')
    added_count   = sum(1 for s in scenes if s.get('type') == '추가씬')

    st.markdown(f"""
    <div class="report-card">
        <h3>11. 각색 원고 (Rewrite Scenes)</h3>
        <div style="background:#FFFBE6;padding:14px;border-radius:8px;
                    border-left:3px solid #FFCB05;margin-bottom:16px;line-height:1.8;">
            <strong style="color:#191970;font-size:0.75rem;display:block;margin-bottom:5px;">✏️ 각색 전략</strong>
            <span style="color:#191970;">{safe(rewrite.get('target_reason', ''))}</span>
        </div>
        <div style="display:flex;gap:10px;margin-bottom:4px;">
            <span style="background:#191970;color:#FFFFFF !important;padding:4px 14px;
                         border-radius:20px;font-size:0.78rem;font-weight:800;">총 {len(scenes)}개 씬</span>
            <span style="background:#2EC484;color:#FFFFFF;padding:4px 14px;
                         border-radius:20px;font-size:0.78rem;font-weight:800;">✏️ 수정씬 {revised_count}개</span>
            <span style="background:#FF6432;color:#FFFFFF;padding:4px 14px;
                         border-radius:20px;font-size:0.78rem;font-weight:800;">✨ 추가씬 {added_count}개</span>
        </div>
    </div>""", unsafe_allow_html=True)

    for i, sc in enumerate(scenes):
        scene_type  = sc.get('type', '수정씬')
        is_revised  = scene_type == '수정씬'
        badge_color = '#2EC484' if is_revised else '#FF6432'
        badge_bg    = '#EDFAF3' if is_revised else '#FFF3EE'
        badge_icon  = '✏️' if is_revised else '✨'
        original    = sc.get('original', '')
        scene_no    = safe(sc.get('scene_no', f'Scene {i+1}'))
        content     = sc.get('content', '')

        insert_between = sc.get('insert_between', '')
        header_extra = f"  ·  📍 {safe(insert_between)}" if (not is_revised and insert_between) else ""

        # 씬 헤더 카드 (expander 대신 div)
        st.markdown(f"""
        <div style="background:#FFFFFF;border:2px solid {badge_color};border-radius:12px;
                    padding:16px 20px;margin-bottom:8px;">
            <div style="font-size:0.88rem;font-weight:900;color:#191970;margin-bottom:10px;">
                {badge_icon} {scene_no}  [{scene_type}]{header_extra}
            </div>""", unsafe_allow_html=True)

        if not is_revised and insert_between:
            st.markdown(f"""
            <div style="background:#EEF0FA;border-left:3px solid #191970;
                        padding:8px 14px;border-radius:6px;margin-bottom:8px;">
                <span style="font-size:0.68rem;font-weight:800;color:#191970;">📍 삽입 위치</span>
                <span style="font-size:0.84rem;color:#191970;margin-left:8px;">{safe(insert_between)}</span>
            </div>""", unsafe_allow_html=True)

        if original:
            st.markdown(f"""
            <div style="background:#F5F5F5;border-left:3px solid #AAAAAA;
                        padding:10px 14px;border-radius:8px;margin-bottom:8px;">
                <div style="font-size:0.68rem;font-weight:800;color:#888;margin-bottom:4px;">📄 기존 씬 (BEFORE)</div>
                <div style="font-size:0.84rem;color:#555;line-height:1.7;">{safe(original)}</div>
            </div>""", unsafe_allow_html=True)

        st.markdown(f"""
        <div style="font-size:0.68rem;font-weight:800;color:{badge_color};margin-bottom:6px;letter-spacing:0.06em;">
            {badge_icon} {scene_type} (AFTER)
        </div>""", unsafe_allow_html=True)

        if content:
            lines_sc = content.replace('\\n', '\n').replace('\n', '\n').split('\n')
            formatted = ""
            for line in lines_sc:
                line_s = line.strip()
                if not line_s:
                    formatted += "<div style='height:8px;'></div>"
                elif line_s.startswith('INT.') or line_s.startswith('EXT.'):
                    formatted += f'<div style="font-weight:900;color:#191970;font-size:0.9rem;letter-spacing:0.05em;margin:10px 0 4px;">{safe(line_s)}</div>'
                elif ':' in line_s and len(line_s.split(':')[0]) < 20 and line_s.split(':')[0].replace(' ','').replace('(','').replace(')','').replace('.','').upper() == line_s.split(':')[0].replace(' ','').replace('(','').replace(')','').replace('.','').upper():
                    parts = line_s.split(':', 1)
                    formatted += f'<div style="margin:6px 0 3px;padding-left:12px;"><span style="font-weight:900;color:#191970;font-size:0.86rem;">{safe(parts[0])}:</span> <span style="color:#333;font-size:0.86rem;line-height:1.7;">{safe(parts[1].strip())}</span></div>'
                else:
                    formatted += f'<div style="color:#444;font-size:0.85rem;line-height:1.75;margin:2px 0;">{safe(line_s)}</div>'

            st.markdown(f"""
            <div style="background:#FAFAFA;border:1px solid #E6E9EF;border-radius:8px;
                        padding:16px 20px;font-family:'Courier Prime','Courier New',monospace;">
                {formatted}
            </div>
        </div>""", unsafe_allow_html=True)
        else:
            st.markdown("</div>", unsafe_allow_html=True)




# =================================================================
# [8] DOCX 생성
# =================================================================
def create_docx(item):
    import subprocess, tempfile, os
    # report_data.json 생성
    report_data = {
        "analysis": item,
        "washing": {
            "washing_table":    item.get("washing_table", []),
            "dialogue_analysis": item.get("dialogue_analysis", {}),
            "suggestions":       item.get("suggestions", [])
        },
        "rewriting": {"rewriting": item.get("rewriting", {})}
    }
    import shutil
    tmp_dir = tempfile.mkdtemp()
    json_path   = os.path.join(tmp_dir, "report_data.json")
    out_path    = os.path.join(tmp_dir, "report_output.docx")

    # gen_docx.js를 tmp_dir에 복사
    script_src = os.path.join(os.path.dirname(os.path.abspath(__file__)), "gen_docx.js")
    if not os.path.exists(script_src):
        for candidate in ["/app/gen_docx.js", "./gen_docx.js", "/home/claude/gen_docx.js"]:
            if os.path.exists(candidate):
                script_src = candidate
                break
    script_path = os.path.join(tmp_dir, "gen_docx.js")
    try:
        shutil.copy2(script_src, script_path)
    except Exception:
        return _create_docx_fallback(item)

    # docx npm 패키지 설치 (Streamlit Cloud에 node_modules 없을 경우)
    nm_path = os.path.join(tmp_dir, "node_modules", "docx")
    if not os.path.exists(nm_path):
        npm_result = subprocess.run(
            ["npm", "install", "docx"],
            capture_output=True, text=True, timeout=120,
            cwd=tmp_dir
        )
        if npm_result.returncode != 0:
            st.warning(f"npm install 실패: {npm_result.stderr[:200]}")
            return _create_docx_fallback(item)

    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(report_data, f, ensure_ascii=False)

    result = subprocess.run(
        ["node", script_path, json_path, out_path],
        capture_output=True, text=True, timeout=60,
        cwd=tmp_dir
    )
    if result.returncode != 0 or not os.path.exists(out_path):
        st.warning(f"DOCX 생성 오류 (Node): {result.stderr[:300]}")
        return _create_docx_fallback(item)

    with open(out_path, "rb") as f:
        data = f.read()

    shutil.rmtree(tmp_dir, ignore_errors=True)
    return data


def _create_docx_fallback(item):
    """Node.js 실패 시 python-docx fallback — 전체 섹션 반영"""

    def safe_int(v, default=0):
        try: return int(float(v))
        except: return default

    def safe_str(v, default=''):
        if v is None: return default
        return str(v)

    try:
        doc = Document()

        # 표지
        t = doc.add_heading("시나리오 검토 보고서", 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"작품명: {safe_str(item.get('title', ''))}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

        # 1. 종합 분석 — 4축 평가표
        doc.add_heading("1. 종합 분석 (Total Analysis) — Hollywood Standard", 1)
        mark = item.get('mark', {}).get('final', 0)
        sc   = item.get('scores', {})
        verdict = item.get('verdict', {})
        doc.add_paragraph(f"최종 점수: {mark} / 10.0   판정: {safe_str(verdict.get('status',''))}")
        doc.add_paragraph(f"판정 근거: {safe_str(verdict.get('rationale',''))}")
        doc.add_heading("4축 정밀 평가 (Hollywood Standard)", 2)
        tbl = doc.add_table(rows=6, cols=4)
        tbl.style = 'Table Grid'
        headers = ['AXIS', '가중치', '평가 기준', '점수']
        for i, h in enumerate(headers):
            tbl.rows[0].cells[i].text = h
        axis_data = [
            ('① STRUCTURE / 구성·플롯',  '30%', '인과관계 정밀도 / 3막 구조 완성도',      safe_int(sc.get('structure', 0))),
            ('② HERO / 캐릭터',          '30%', 'Goal / Need / Strategy / 감정선',          safe_int(sc.get('hero', 0))),
            ('③ CONCEPT / 소재·컨셉',    '20%', '하이컨셉 · 독창성 · 시장성',              safe_int(sc.get('concept', 0))),
            ('④ GENRE / 장르 적합성',    '20%', '장르 문법 충실도 / 타깃 소구력',           safe_int(sc.get('genre', 0))),
        ]
        for i, (axis, weight, criteria, val) in enumerate(axis_data, 1):
            row = tbl.rows[i].cells
            row[0].text = axis
            row[1].text = weight
            row[2].text = criteria
            clamped = min(max(val, 0), 10)
            row[3].text = f"{'█'*clamped}{'░'*(10-clamped)}  {val}/10"
        final_row = tbl.rows[5].cells
        final_row[0].text = 'FINAL'
        final_row[1].text = ''
        final_row[2].text = '0.3S + 0.3H + 0.2C + 0.2G'
        final_row[3].text = f"{mark} / 10.0"

        # 2. 로그라인
        doc.add_heading("2. 로그라인 분석 (Logline Pack)", 1)
        log = item.get('logline', {})
        if isinstance(log, dict):
            doc.add_heading("ORIGINAL", 2)
            doc.add_paragraph(safe_str(log.get('original', '')))
            doc.add_heading("WASHED", 2)
            doc.add_paragraph(safe_str(log.get('washed', '')))
        else:
            doc.add_paragraph(safe_str(log))

        # 3. 줄거리
        doc.add_heading("3. 줄거리 (Synopsis)", 1)
        doc.add_paragraph(safe_str(item.get('synopsis', '')))

        # 4. 장단점 + 핵심처방
        doc.add_heading("4. 장점 및 보완점 (Pros & Cons)", 1)
        pc = item.get('pros_cons', {})
        if isinstance(pc, dict):
            for p in pc.get('pros', []):  doc.add_paragraph(f"✅ {safe_str(p)}")
            for c in pc.get('cons', []):  doc.add_paragraph(f"⚠️ {safe_str(c)}")
            if pc.get('prescription'):
                doc.add_heading("핵심 처방 (Key Prescription)", 2)
                doc.add_paragraph(safe_str(pc.get('prescription', '')))

        # 5. 서사 동력
        doc.add_heading("5. 서사 동력 (Narrative Drive)", 1)
        drive = item.get('drive', {})
        if isinstance(drive, dict):
            doc.add_paragraph(f"① 목적(욕망): {safe_str(drive.get('goal',''))}")
            doc.add_paragraph(f"② 발생요인: {safe_str(drive.get('lack',''))}")
            doc.add_paragraph(f"③ 해결전략: {safe_str(drive.get('strategy',''))}")
            ev = drive.get('evaluation', {})
            if isinstance(ev, dict):
                doc.add_paragraph(f"목적 명확성: {safe_str(ev.get('clarity',''))}")
                doc.add_paragraph(f"발생요인 확실성: {safe_str(ev.get('urgency',''))}")
                doc.add_paragraph(f"해결전략 창의성: {safe_str(ev.get('consistency',''))}")
                doc.add_paragraph(safe_str(ev.get('overall_diagnosis', '')))

        # 6. 15-Beat Sheet
        BEAT_KO = {
            "Opening Image":"오프닝 이미지","Theme Stated":"주제 제시","Set-up":"설정",
            "Catalyst":"촉매","Debate":"갈등","Break Into Two":"2막 진입",
            "B Story":"B 스토리","Fun and Games":"재미와 게임","Midpoint":"중간점",
            "Bad Guys Close In":"위기 고조","All Is Lost":"모든 것을 잃다",
            "Dark Night of Soul":"영혼의 어둔 밤","Dark Night of the Soul":"영혼의 어둔 밤",
            "Break Into Three":"3막 진입","Finale":"피날레","Climax":"클라이맥스",
            "Final Image":"최종 이미지",
        }
        circles_fb = ['①','②','③','④','⑤','⑥','⑦','⑧','⑨','⑩','⑪','⑫','⑬','⑭','⑮']
        doc.add_heading("6. 구성 및 플롯 (15-Beat Sheet)", 1)
        beats = item.get('beats', {})
        if isinstance(beats, dict) and beats:
            btbl = doc.add_table(rows=1, cols=2)
            btbl.style = 'Table Grid'
            btbl.rows[0].cells[0].text = 'BEAT'
            btbl.rows[0].cells[1].text = 'DESCRIPTION'
            for idx, (k, v) in enumerate(sorted(beats.items()), 0):
                name = re.sub(r'^[\d\.\-\_\s]+', '', str(k)).strip()
                ko   = BEAT_KO.get(name, name)
                row  = btbl.add_row()
                row.cells[0].text = f"{circles_fb[idx] if idx < 15 else ''} {ko} ({name})"
                row.cells[1].text = safe_str(v)

        # 7. 장르
        doc.add_heading("7. 장르 분석 및 적합도 (Genre Compliance)", 1)
        genre = item.get('genre_compliance', item.get('genre_suitability', {}))
        if isinstance(genre, dict):
            doc.add_paragraph(f"장르: {safe_str(genre.get('genre_key', genre.get('genre_name', '')))}")
            doc.add_paragraph(f"준수도: {genre.get('compliance_score',0)} / 10")
            # 장르적 재미 진단
            genre_fun_alive = genre.get('genre_fun_alive', None)
            if genre_fun_alive is not None:
                doc.add_paragraph(f"장르적 재미: {'✓ 작동' if genre_fun_alive else '✗ 약함'}")
            if genre.get('genre_fun_diagnosis'):
                doc.add_paragraph(f"장르적 재미 진단: {safe_str(genre.get('genre_fun_diagnosis',''))}")
            # must_have_check
            must_checks = genre.get('must_have_check', [])
            if isinstance(must_checks, list) and must_checks:
                doc.add_heading("필수 요소 체크", 2)
                for c in must_checks:
                    if isinstance(c, dict):
                        status_icon = '✅' if c.get('status') == '충족' else ('△' if c.get('status') == '약함' else '❌')
                        doc.add_paragraph(f"{status_icon} {safe_str(c.get('item',''))} [{safe_str(c.get('status',''))}] — {safe_str(c.get('evidence',''))}")
            else:
                for c in genre.get('checks', []):
                    doc.add_paragraph(f"✅ {safe_str(c)}")
            # Hook/Punch 체크
            hp = genre.get('hook_punch_check', {})
            if isinstance(hp, dict) and hp:
                doc.add_heading("Hook / Punch 체크", 2)
                doc.add_paragraph(f"{'✓' if hp.get('hook_present') else '✗'} Hook: {safe_str(hp.get('hook_note',''))}")
                doc.add_paragraph(f"{'✓' if hp.get('punch_present') else '✗'} Punch: {safe_str(hp.get('punch_note',''))}")
            # 실패 패턴
            fail_patterns = genre.get('fail_patterns_found', [])
            if isinstance(fail_patterns, list) and fail_patterns:
                doc.add_heading("발견된 실패 패턴", 2)
                for f in fail_patterns:
                    doc.add_paragraph(f"⚠️ {safe_str(f)}")
            for m in genre.get('missing_elements', []):
                doc.add_paragraph(f"❌ {safe_str(m)}")
            if genre.get('doctoring'):
                doc.add_paragraph(safe_str(genre.get('doctoring', '')))

        # 8. 시퀀스 워싱
        doc.add_page_break()
        doc.add_heading("8. 시퀀스 워싱 (Washing Table)", 1)
        for row in item.get('washing_table', []):
            if not isinstance(row, dict): continue
            doc.add_paragraph(f"[{safe_str(row.get('seq',''))}]  {safe_str(row.get('label',''))}")
            wtbl = doc.add_table(rows=1, cols=2)
            wtbl.style = 'Table Grid'
            wtbl.rows[0].cells[0].text = f"⚠️ 진단\n{safe_str(row.get('diagnosis',''))}"
            wtbl.rows[0].cells[1].text = f"✅ 처방\n{safe_str(row.get('prescription',''))}"

        # 9. 대사 워싱
        doc.add_heading("9. 대사 워싱 (Dialogue Washing)", 1)
        da = item.get('dialogue_analysis', {})
        if isinstance(da, dict) and da:
            doc.add_paragraph(f"종합 대사 수준: {da.get('overall_score',0)} / 10")
            ax = da.get('axis_scores', {})
            if isinstance(ax, dict) and ax:
                doc.add_heading("3축 점수", 2)
                atbl = doc.add_table(rows=4, cols=3)
                atbl.style = 'Table Grid'
                atbl.rows[0].cells[0].text = '평가 축'
                atbl.rows[0].cells[1].text = '기준'
                atbl.rows[0].cells[2].text = '점수'
                axes = [
                    ('① 캐릭터 적합성', '고유 어휘·말투·감정 반영',   safe_int(ax.get('character_voice',0))),
                    ('② 서브텍스트',    '표면↔이면 충돌, 설명형 금지', safe_int(ax.get('subtext',0))),
                    ('③ 행동/감정/관계','장면 추진력, 정보전달 금지',   safe_int(ax.get('action_driven',0))),
                ]
                for i, (lbl, crit, val) in enumerate(axes, 1):
                    atbl.rows[i].cells[0].text = lbl
                    atbl.rows[i].cells[1].text = crit
                    clamped = min(max(val, 0), 10)
                    atbl.rows[i].cells[2].text = f"{'█'*clamped}{'░'*(10-clamped)}  {val}/10"
            doc.add_paragraph(safe_str(da.get('overall_comment', '')))
            for s in da.get('strengths', []):
                doc.add_paragraph(f"💪 {safe_str(s)}")
            issues = da.get('issues', [])
            if isinstance(issues, list) and issues:
                doc.add_heading("대사 4축 진단 Before / After", 2)
                for issue in issues:
                    if not isinstance(issue, dict): continue
                    doc.add_paragraph(f"[{safe_str(issue.get('type',''))}]  {safe_str(issue.get('axis',''))}  {safe_str(issue.get('description',''))}")
                    itbl = doc.add_table(rows=1, cols=2)
                    itbl.style = 'Table Grid'
                    itbl.rows[0].cells[0].text = f"❌ BEFORE\n{safe_str(issue.get('example_bad',''))}"
                    itbl.rows[0].cells[1].text = f"✅ ④ 개선 제안\n{safe_str(issue.get('example_good',''))}"
                    if issue.get('rewrite_note'):
                        doc.add_paragraph(f"✏️ Moon 지시: {safe_str(issue.get('rewrite_note',''))}")
        
        # 10. 각색 제안
        doc.add_heading("10. 각색 제안 (Action Plan)", 1)
        for i, s in enumerate(item.get('suggestions', []), 1):
            clean_s = re.sub(r'^[\d\.\s]+', '', safe_str(s)).strip()
            doc.add_paragraph(f"STEP {i:02d}  {clean_s}")

        # 11. 각색 원고
        doc.add_page_break()
        doc.add_heading("11. 각색 원고 (Rewrite Scenes)", 1)
        rw = item.get('rewriting', {})
        if isinstance(rw, dict):
            if rw.get('target_reason'):
                doc.add_paragraph(f"✏️ 각색 전략: {safe_str(rw.get('target_reason',''))}")
            for sc in rw.get('scenes', []):
                if not isinstance(sc, dict): continue
                is_r = sc.get('type') == '수정씬'
                doc.add_heading(f"{safe_str(sc.get('scene_no',''))}  [{'✏️ 수정씬' if is_r else '✨ 추가씬'}]", 2)
                if sc.get('original'):
                    doc.add_paragraph(f"📄 기존 씬 (BEFORE)\n{safe_str(sc.get('original',''))}")
                content = safe_str(sc.get('content', ''))
                content = content.replace('\\n', '\n')
                doc.add_paragraph(f"{'✏️ 수정씬' if is_r else '✨ 추가씬'} (AFTER)\n{content}")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    except Exception as e:
        st.error(f"DOCX 생성 중 오류 발생: {type(e).__name__} — {e}")
        # 최소한의 오류 보고서라도 생성
        try:
            doc_err = Document()
            doc_err.add_heading("DOCX 생성 오류", 0)
            doc_err.add_paragraph(f"오류: {type(e).__name__} — {e}")
            doc_err.add_paragraph(f"작품명: {safe_str(item.get('title', ''))}")
            doc_err.add_paragraph("전체 보고서 생성에 실패했습니다. Streamlit 화면의 분석 결과를 참고해주세요.")
            buf_err = io.BytesIO()
            doc_err.save(buf_err)
            buf_err.seek(0)
            return buf_err.getvalue()
        except:
            return None


# =================================================================
# [9] 진행 바
# =================================================================
def render_step_bar(step):
    labels = ["📄 업로드", "🔍 CHRIS", "🧹 SHIHO", "✍️ MOON"]
    dots = ""
    for i, label in enumerate(labels):
        cls = "done" if i < step else ("active" if i == step else "wait")
        sym = "✓" if i < step else str(i + 1)
        dots += (f'<div style="display:flex;flex-direction:column;align-items:center;gap:5px;">'
                 f'<div class="step-dot {cls}">{sym}</div>'
                 f'<div style="font-size:0.68rem;opacity:0.65;white-space:nowrap;color:#191970;">{label}</div></div>')
        if i < len(labels) - 1:
            lc = "done" if i < step - 1 else ""
            dots += f'<div class="step-line {lc}" style="margin-bottom:20px;"></div>'
    st.markdown(f'<div class="step-bar">{dots}</div>', unsafe_allow_html=True)

# =================================================================
# [10] 비서 카드
# =================================================================
AVATAR_MAP = {
    "CHRIS": CHRIS_SVG,
    "SHIHO": SHIHO_SVG,
    "MOON":  MOON_SVG,
}

def agent_card(emoji, name, role, desc, status, key, btn_label, btn_fn, result_fn, result_data):
    badge_map = {
        'waiting': '<span style="background:#E6E9EF;color:#888888;padding:2px 10px;border-radius:20px;font-size:0.72rem;font-weight:800;">대기 중</span>',
        'active':  '<span style="background:#FFCB05;color:#191970;padding:2px 10px;border-radius:20px;font-size:0.72rem;font-weight:800;">● 준비됨</span>',
        'done':    '<span style="background:#2EC484;color:#FFFFFF;padding:2px 10px;border-radius:20px;font-size:0.72rem;font-weight:800;">✓ 완료</span>',
    }
    avatar_html = AVATAR_MAP.get(name, f'<div style="font-size:2.2rem;">{emoji}</div>')
    st.markdown(f"""
    <div class="agent-card {status}">
        <div style="display:flex;align-items:flex-start;gap:18px;">
            <div style="width:72px;height:72px;flex-shrink:0;border-radius:16px;overflow:hidden;">{avatar_html}</div>
            <div style="flex:1;">
                <div style="display:flex;align-items:center;gap:10px;margin-bottom:3px;">
                    <span style="font-size:1.15rem;font-weight:900;letter-spacing:0.04em;color:#191970;">{name}</span>
                    {badge_map[status]}
                </div>
                <div style="font-size:0.75rem;color:#B8860B;font-weight:700;margin-bottom:4px;letter-spacing:0.06em;">{role}</div>
                <div style="font-size:0.83rem;color:#191970;opacity:0.65;line-height:1.5;">{desc}</div>
            </div>
        </div>
    </div>""", unsafe_allow_html=True)

    if status == 'active':
        if st.button(btn_label, key=f"{key}_btn", use_container_width=True):
            btn_fn()

    if status == 'done' and result_data:
        result_fn(result_data)

# =================================================================
# [11] 워크스페이스 페이지
# =================================================================
def show_workspace():
    st.markdown("""
    <div style="text-align:center;padding:36px 0 8px;">
        <div style="font-size:0.72rem;font-weight:900;letter-spacing:0.5em;color:#FFCB05;margin-bottom:6px;">BLUE JEANS PICTURES</div>
        <div style="font-size:2.8rem;font-weight:950;color:#191970;line-height:1;letter-spacing:-0.03em;">REWRITE ENGINE</div>
        <div style="margin-top:6px;font-size:0.75rem;letter-spacing:0.15em;color:#191970;opacity:0.4;">YOUNG · VINTAGE · FREE · INNOVATIVE</div>
    </div>""", unsafe_allow_html=True)

    step = st.session_state.step
    render_step_bar(step)
    st.markdown('<hr>', unsafe_allow_html=True)

    client = get_client()

    # ── 업로드 ──
    if step == 0:


        # 파일 업로드 전: uploader + 커스텀 드롭존 안내
        # 파일 업로드 후: uploader 숨기고 파일 정보 카드만 표시
        if not st.session_state.get('_uploaded_ready'):
            uploaded = st.file_uploader(
                "분석할 시나리오 PDF를 업로드하세요",
                type=["pdf"],
                key="pdf_uploader"
            )
            if not uploaded:
                pass
            else:
                with st.spinner("📄 텍스트 추출 중..."):
                    text = extract_text(uploaded)
                if text:
                    st.session_state.raw_text = text
                    st.session_state['_uploaded_name'] = uploaded.name
                    st.session_state['_uploaded_kb']   = round(len(uploaded.getvalue()) / 1024)
                    st.session_state['_uploaded_chars'] = len(text)
                    st.session_state['_uploaded_ready'] = True
                    st.rerun()
        else:
            # 업로드 완료 → uploader 숨기고 파일 카드만 표시
            fname  = st.session_state.get('_uploaded_name', '')
            kb     = st.session_state.get('_uploaded_kb', 0)
            chars  = st.session_state.get('_uploaded_chars', 0)
            st.markdown(f"""
            <div style="background:#EEF0FA;border:1px solid #C5D0F0;border-radius:12px;
                        padding:16px 20px;margin-bottom:16px;display:flex;align-items:center;gap:14px;">
                <div style="font-size:2rem;">📄</div>
                <div style="flex:1;">
                    <div style="font-weight:800;color:#191970;font-size:0.95rem;">{safe(fname)}</div>
                    <div style="font-size:0.8rem;color:#666;margin-top:3px;">{kb:,} KB · {chars:,}자 추출 완료</div>
                </div>
                <div style="background:#2EC484;color:#FFFFFF;padding:4px 12px;border-radius:20px;
                            font-size:0.75rem;font-weight:800;">✓ 준비완료</div>
            </div>
            """, unsafe_allow_html=True)
            col1, col2, col3 = st.columns([2, 1, 2])
            with col2:
                if st.button("🚀 분석 시작", use_container_width=True):
                    st.session_state.step = 1
                    st.session_state['_uploaded_ready'] = False
                    st.rerun()
            # 파일 다시 선택 링크
            st.markdown('<div style="text-align:center;margin-top:8px;">', unsafe_allow_html=True)
            if st.button("↩ 다른 파일 선택", use_container_width=False):
                for k in ['_uploaded_ready','_uploaded_name','_uploaded_kb','_uploaded_chars']:
                    st.session_state.pop(k, None)
                st.session_state.raw_text = None
                st.rerun()
            st.markdown('</div>', unsafe_allow_html=True)
        return

    text = st.session_state.raw_text

    # ── BLUE 비서 ──
    def do_chris():
        with st.spinner("🔍 Chris가 시나리오를 분석하고 있습니다... (약 30~60초 소요)"):
            # API 키 확인
            api_key = st.secrets.get("ANTHROPIC_API_KEY", "")
            if not api_key:
                st.error("❌ ANTHROPIC_API_KEY가 Secrets에 없습니다. Streamlit Cloud → Settings → Secrets 확인")
                return
            r = run_blue(text, client)
            if r:
                st.session_state.analysis = r
                st.session_state.step = 2
                st.rerun()
            else:
                st.error("분석 실패. 위의 API 오류 메시지를 확인하세요.")

    chris_st = 'done' if st.session_state.analysis else 'active'
    agent_card("🔍", "CHRIS", "Senior Script Analyst",
               "구조 분석 · 15-Beat Sheet · 서사동력 · 시각화 · 장르 진단",
               chris_st, "chris", "🔍 Chris 분석 시작", do_chris,
               render_analysis, st.session_state.analysis)

    if not st.session_state.analysis:
        agent_card("🧹", "SHIHO", "Script Doctor",
                   "시퀀스 워싱 · 문제 유형 진단 · 처방 + Risk · 각색 제안 10가지",
                   'waiting', "shiho", "", lambda: None, lambda x: None, None)
        agent_card("✍️", "MOON", "Senior Screenwriter",
                   "표준 각본 형식 · 10개 씬 리라이팅",
                   'waiting', "moon", "", lambda: None, lambda x: None, None)
        return

    # ── JEAN 비서 ──
    def do_shiho():
        with st.spinner("🧹 Shiho가 시퀀스를 워싱하고 있습니다... (약 30~60초 소요)"):
            r = run_jean(text, st.session_state.analysis, client)
            if r:
                st.session_state.washing = r
                st.session_state.step = 3
                st.rerun()
            else:
                st.error("워싱 실패 — API 오류 메시지를 위에서 확인하세요.")

    shiho_st = 'done' if st.session_state.washing else 'active'
    agent_card("🧹", "SHIHO", "Script Doctor",
               "시퀀스 워싱 · 문제 유형 진단 · 처방 + Risk · 각색 제안 10가지",
               shiho_st, "shiho", "🧹 Shiho 워싱 시작", do_shiho,
               render_washing, st.session_state.washing)

    if not st.session_state.washing:
        agent_card("✍️", "MOON", "Senior Screenwriter",
                   "표준 각본 형식 · 10개 씬 리라이팅",
                   'waiting', "moon", "", lambda: None, lambda x: None, None)
        return

    # ── PICTURES 비서 ──
    def do_moon():
        with st.spinner("✍️ Moon이 각색 원고를 작성하고 있습니다... (약 60~90초 소요)"):
            r = run_pictures(text, st.session_state.washing, client)
            if r:
                st.session_state.rewriting = r
                full = {
                    **st.session_state.analysis,
                    'washing_table':     st.session_state.washing.get('washing_table', []),
                    'suggestions':       st.session_state.washing.get('suggestions', []),
                    'dialogue_analysis': st.session_state.washing.get('dialogue_analysis', {}),
                    'rewriting':         r.get('rewriting', {})
                }
                st.session_state.db.insert(0, full)
                st.session_state.selected_item = full
                st.session_state.step = 4
                st.rerun()
            else:
                st.error("리라이팅 실패. 위의 API 오류 메시지를 확인하세요.")

    moon_st = 'done' if st.session_state.rewriting else 'active'
    agent_card("✍️", "MOON", "Senior Screenwriter",
               "오프닝 훅 / 치명상 구간 선정 · 표준 각본 형식 · 10개 씬 리라이팅",
               moon_st, "moon", "✍️ Moon 리라이팅 시작", do_moon,
               render_rewriting, st.session_state.rewriting)

    # ── 완료 ──
    if st.session_state.step == 4 and st.session_state.selected_item:
        st.markdown('<hr>', unsafe_allow_html=True)
        st.markdown('<div style="text-align:center;padding:16px 0;"><span style="font-size:1.4rem;font-weight:900;color:#2EC484;">🎉 모든 분석 완료!</span></div>', unsafe_allow_html=True)
        item = st.session_state.selected_item
        c1, c2, c3 = st.columns(3)
        with c1:
            title = re.sub(r'[/*?:"<>|]', '_', item.get('title', '제목없음'))
            try:
                docx_bytes = create_docx(item)
                st.download_button(
                    "📄 보고서 다운로드 (DOCX)",
                    data=docx_bytes,
                    file_name=f"시나리오검토보고서_{title}_Blue.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="btn_download_docx",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"DOCX 생성 오류: {e}")
        with c2:
            if st.button("🏠 목록으로 돌아가기", key="btn_to_index", use_container_width=True):
                st.session_state.page = "index"
                st.rerun()
        with c3:
            if st.button("🔄 새 시나리오 분석", key="btn_new_analysis", use_container_width=True):
                for k in ['step', 'raw_text', 'analysis', 'washing', 'rewriting', 'selected_item']:
                    st.session_state[k] = 0 if k == 'step' else None
                st.rerun()

# =================================================================
# [12] 갤러리 페이지 (홈)
# =================================================================
def show_index():
    st.markdown("""
    <div style="text-align:center;padding:60px 0 36px;">
        <div style="font-size:0.72rem;font-weight:900;letter-spacing:0.5em;color:#FFCB05;">BLUE JEANS PICTURES</div>
        <div style="font-size:3.6rem;font-weight:950;color:#191970;line-height:1;letter-spacing:-0.03em;margin:10px 0;">REWRITE ENGINE</div>
        <div style="font-size:0.78rem;letter-spacing:0.2em;color:#191970;opacity:0.35;">YOUNG · VINTAGE · FREE · INNOVATIVE</div>
    </div>""", unsafe_allow_html=True)

    _, c, _ = st.columns([1, 1, 1])
    with c:
        if st.button("+ 새 시나리오 분석하기", use_container_width=True):
            for k in ['step', 'raw_text', 'analysis', 'washing', 'rewriting']:
                st.session_state[k] = 0 if k == 'step' else None
            st.session_state.page = "workspace"
            st.rerun()

    st.markdown('<hr>', unsafe_allow_html=True)
    st.markdown("### 📋 검토 보고서 아카이브")

    if not st.session_state.db:
        st.markdown(
            '<div style="text-align:center;padding:60px 0;opacity:0.3;">'
            '<div style="font-size:3rem;margin-bottom:10px;">🎬</div>'
            '<div style="color:#191970;">아직 분석된 시나리오가 없습니다.</div></div>',
            unsafe_allow_html=True
        )
    else:
        cols = st.columns(3)
        for idx, item in enumerate(st.session_state.db):
            with cols[idx % 3]:
                img_url = get_genre_img_url(item)
                genre_label = item.get('genre', {}).get('primary', '') if isinstance(item.get('genre'), dict) else str(item.get('genre', ''))
                score_val = item.get('mark', {}).get('final', 0)
                score_color = '#2EC484' if float(score_val) >= 7 else ('#FFCB05' if float(score_val) >= 5 else '#FF6432')
                st.markdown(f"""
                <div class="gallery-card">
                    <img class="gallery-card-img" src="{img_url}"
                         onerror="this.style.display='none';this.nextElementSibling.style.display='flex';"
                         alt="{safe(item.get('title',''))}">
                    <div style="display:none;height:140px;background:linear-gradient(135deg,#191970,#3A3A8C);
                                align-items:center;justify-content:center;font-size:2.8rem;">🎬</div>
                    <div class="gallery-card-body">
                        <div style="font-size:0.65rem;font-weight:700;color:#888;letter-spacing:0.1em;
                                    text-transform:uppercase;margin-bottom:6px;">{safe(genre_label)}</div>
                        <div style="font-weight:900;font-size:1rem;margin-bottom:10px;color:#191970;
                                    line-height:1.3;">{safe(item.get('title', 'Untitled'))}</div>
                        <div style="display:flex;align-items:center;justify-content:center;gap:6px;">
                            <div style="font-size:1.6rem;font-weight:950;color:{score_color};">{score_val}</div>
                            <div style="text-align:left;">
                                <div style="font-size:0.65rem;color:#888;">/10.0</div>
                                <div style="font-size:0.6rem;color:{score_color};font-weight:800;">
                                    {"우수" if float(score_val) >= 8 else ("양호" if float(score_val) >= 6 else ("보통" if float(score_val) >= 4 else "미흡"))}
                                </div>
                            </div>
                        </div>
                    </div>
                </div>""", unsafe_allow_html=True)
                c1, c2 = st.columns(2)
                with c1:
                    if st.button("보기", key=f"v_{idx}", use_container_width=True):
                        st.session_state.selected_item = item
                        st.session_state.analysis  = item
                        st.session_state.washing   = {
                            'washing_table':    item.get('washing_table', []),
                            'dialogue_analysis': item.get('dialogue_analysis', {}),
                            'suggestions':       item.get('suggestions', [])
                        }
                        st.session_state.rewriting = {'rewriting': item.get('rewriting', {})}
                        st.session_state.step = 4
                        st.session_state.page = "workspace"
                        st.rerun()
                with c2:
                    if st.button("삭제", key=f"d_{idx}", use_container_width=True):
                        st.session_state.db.pop(idx)
                        st.rerun()

# =================================================================
# [13] 실행
# =================================================================
apply_design()

if st.session_state.page == "index":
    show_index()
elif st.session_state.page == "workspace":
    show_workspace()

st.markdown(
    '<div style="text-align:center;font-size:0.65rem;padding:40px 0 16px;'
    'letter-spacing:2px;color:#191970;opacity:0.25;">© 2026 BLUE JEANS PICTURES</div>',
    unsafe_allow_html=True
)
